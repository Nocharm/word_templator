"""20페이지 비정규화 docx 라운드트립 검증 — Phase 1~4 전체 파이프라인 점검."""

import io
import json
import sys
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# Ensure repo root on path
ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from app.domain.style_spec import StyleSpec  # noqa: E402
from app.parser.parse_docx import parse_docx  # noqa: E402
from app.renderer.render_docx import render_docx  # noqa: E402
from tests.fixtures.build_full_messy_sample import build_full_messy_sample  # noqa: E402


def _load_default_spec() -> StyleSpec:
    seed = ROOT / "app" / "templates_seed" / "report.json"
    data = json.loads(seed.read_text(encoding="utf-8"))
    return StyleSpec.model_validate(data["spec"])


def section(title: str) -> None:
    print(f"\n=== {title} ===")


def check(label: str, ok: bool, detail: str = "") -> bool:
    icon = "✅" if ok else "❌"
    sfx = f" — {detail}" if detail else ""
    print(f"  {icon} {label}{sfx}")
    return ok


def main() -> int:
    import os
    import tempfile

    bugs: list[str] = []

    with tempfile.TemporaryDirectory() as tmp:
        os.environ["DATA_DIR"] = tmp
        src = Path(tmp) / "full.docx"
        build_full_messy_sample(src)

        section("픽스처")
        size_kb = src.stat().st_size / 1024
        print(f"  생성: {src} ({size_kb:.1f} KB)")

        # ── PARSE ──
        user_id = uuid.uuid4()
        job_id = uuid.uuid4()
        outline = parse_docx(
            src.read_bytes(),
            filename="full.docx",
            user_id=user_id,
            job_id=job_id,
        )
        blocks = outline.blocks

        section("PARSE 결과 통계")
        kinds = {
            k: sum(1 for b in blocks if b.kind == k)
            for k in ("paragraph", "table", "image", "field")
        }
        print(f"  총 블록: {len(blocks)}")
        print(
            f"  kind 분포: paragraph={kinds['paragraph']} table={kinds['table']} "
            f"image={kinds['image']} field={kinds['field']}"
        )

        levels = {}
        for b in blocks:
            if b.kind == "paragraph":
                levels[b.level] = levels.get(b.level, 0) + 1
        print(f"  레벨 분포: {dict(sorted(levels.items()))}")

        detected_by = {}
        for b in blocks:
            if b.kind == "paragraph" and b.level >= 1:
                detected_by[b.detected_by] = detected_by.get(b.detected_by, 0) + 1
        print(f"  헤딩 감지 출처: {detected_by}")

        # ── 검증: 헤딩 ──
        section("검증 #1 — 헤딩 감지")
        h1_count = sum(1 for b in blocks if b.kind == "paragraph" and b.level == 1)
        h2_count = sum(1 for b in blocks if b.kind == "paragraph" and b.level == 2)
        h3_count = sum(1 for b in blocks if b.kind == "paragraph" and b.level == 3)
        # Expected: 5 main H1 + 표지 휴리스틱 + 부록 *** ***
        # Expected: ≥ 3 H2, ≥ 2 H3 (2.2.1 / 2.2.2)
        if not check("H1 ≥ 5개 감지", h1_count >= 5, f"실제 {h1_count}개"):
            bugs.append(f"H1 부족: {h1_count} < 5")
        if not check("H2 ≥ 3개 감지", h2_count >= 3, f"실제 {h2_count}개"):
            bugs.append(f"H2 부족: {h2_count} < 3")
        if not check("H3 ≥ 2개 감지 (2.2.1, 2.2.2)", h3_count >= 2, f"실제 {h3_count}개"):
            bugs.append(f"H3 부족: {h3_count} < 2")
        # Cover-page heuristic should mark "2026 사내 시스템 표준화 결과 보고서" as heading
        cover = next((b for b in blocks if b.text and "2026 사내 시스템" in b.text), None)
        if cover is None:
            bugs.append("표지 제목이 outline 에 없음")
            check("표지 제목 outline 포함", False)
        else:
            check("표지 제목 outline 포함", cover.level >= 1, f"level={cover.level}")
            if cover.level == 0:
                bugs.append(f"표지 제목 휴리스틱 실패: level={cover.level}")

        # 부록 *** *** 휴리스틱
        appendix = [
            b
            for b in blocks
            if b.kind == "paragraph" and b.text and "부록" in b.text and "***" in b.text
        ]
        if not check(
            "*** 부록 *** 휴리스틱 헤딩 인식",
            all(b.level >= 1 for b in appendix),
            f"{[b.level for b in appendix]}",
        ):
            bugs.append(f"부록 *** *** 휴리스틱 실패: 레벨 {[b.level for b in appendix]}")

        # ── 검증: 표 ──
        section("검증 #2 — 표")
        tables = [b for b in blocks if b.kind == "table"]
        if not check("표 2개 추출", len(tables) == 2, f"실제 {len(tables)}개"):
            bugs.append(f"표 개수 미스매치: {len(tables)} ≠ 2")
        for i, t in enumerate(tables):
            if not check(
                f"표[{i}] markdown + raw_ref 채워짐",
                bool(t.markdown) and bool(t.raw_ref),
                f"md={'O' if t.markdown else 'X'}, raw_ref={t.raw_ref}",
            ):
                bugs.append(f"표[{i}] markdown 또는 raw_ref 비어있음")
        # 캡션 흡수: "표 1. 부서별 의견 분류"는 표 위에, "표 2. 변환 전후 비교"는 표 아래에
        cap1 = next((t for t in tables if t.caption and "부서별" in t.caption), None)
        cap2 = next((t for t in tables if t.caption and "변환" in t.caption), None)
        if not check("표 1 캡션(위) 흡수", cap1 is not None, str(cap1.caption if cap1 else None)):
            bugs.append("표 1 캡션 (표 위) 미흡수")
        if not check("표 2 캡션(아래) 흡수", cap2 is not None, str(cap2.caption if cap2 else None)):
            bugs.append("표 2 캡션 (표 아래) 미흡수")
        # 흡수된 캡션 paragraph 가 outline 에 남아있으면 안됨
        leftover = [
            b for b in blocks if b.kind == "paragraph" and b.text and "표 1. 부서별" in b.text
        ]
        if not check(
            "흡수된 표1 캡션 paragraph 제거됨", len(leftover) == 0, f"잔존 {len(leftover)}"
        ):
            bugs.append("표 1 캡션 paragraph 가 outline 에 중복 존재")

        # 병합셀 마커
        merged_marker = any(t.markdown and "[병합셀]" in t.markdown for t in tables)
        if not check("표 2 병합셀 마커 [병합셀] 표시", merged_marker):
            bugs.append("표 2 병합셀 마커 누락")

        # ── 검증: 이미지 ──
        section("검증 #3 — 이미지")
        images = [b for b in blocks if b.kind == "image"]
        if not check("이미지 1개 추출", len(images) == 1, f"실제 {len(images)}개"):
            bugs.append(f"이미지 개수 미스매치: {len(images)} ≠ 1")
        if images:
            img = images[0]
            if not check(
                "이미지 raw_ref + preview_url 채워짐",
                bool(img.raw_ref) and bool(img.preview_url),
                f"raw_ref={img.raw_ref}, preview_url={img.preview_url}",
            ):
                bugs.append("이미지 raw_ref/preview_url 미설정")
            if not check(
                "이미지 캡션(위) 흡수",
                bool(img.caption and "그림 1" in img.caption),
                str(img.caption),
            ):
                bugs.append("이미지 캡션 (그림 1) 미흡수")

        # 디스크 보존
        img_dir = Path(tmp) / "images" / str(job_id)
        if img_dir.exists():
            saved = list(img_dir.glob("image-*"))
            check(f"이미지 디스크 저장 ({len(saved)}개)", len(saved) >= 1)
        else:
            bugs.append("이미지 디스크 폴더 없음")
            check("이미지 디스크 저장", False, "image_dir 누락")

        # ── 검증: 필드/북마크 ──
        section("검증 #4 — 필드 / 북마크")
        fields = [b for b in blocks if b.kind == "paragraph" and b.raw_xml_ref]
        # Expected: TOC + (REF body 1개) + (REF/PAGEREF body 1개) + 4개 북마크 헤딩 = 7개 정도
        if not check(
            "필드/북마크 보존된 paragraph ≥ 4개", len(fields) >= 4, f"실제 {len(fields)}개"
        ):
            bugs.append(f"필드/북마크 보존 부족: {len(fields)} < 4")

        toc_blocks = [b for b in fields if b.field_kind == "toc"]
        ref_blocks = [b for b in fields if b.field_kind == "ref"]
        pageref_blocks = [b for b in fields if b.field_kind == "pageref"]
        if not check("TOC 1개 분류", len(toc_blocks) == 1, f"실제 {len(toc_blocks)}"):
            bugs.append(f"TOC 분류 오류: {len(toc_blocks)} ≠ 1")
        # 픽스처: 본문 1 = REF only, 본문 2 = REF + PAGEREF (PAGEREF 우선순위로 분류).
        if not check(
            "REF ≥ 1개 분류 (REF only paragraph)", len(ref_blocks) >= 1, f"실제 {len(ref_blocks)}"
        ):
            bugs.append(f"REF 부족: {len(ref_blocks)} < 1")
        if not check(
            "PAGEREF ≥ 1개 분류 (REF+PAGEREF mixed paragraph)",
            len(pageref_blocks) >= 1,
            f"실제 {len(pageref_blocks)}",
        ):
            bugs.append(f"PAGEREF 부족: {len(pageref_blocks)} < 1")

        # 디스크
        raw_dir = Path(tmp) / "docs" / str(user_id) / str(job_id) / "raw"
        if raw_dir.exists():
            field_files = list(raw_dir.glob("field-*.xml"))
            table_files = list(raw_dir.glob("table-*.xml"))
            check(f"필드 paragraph 디스크 저장 ({len(field_files)}개)", len(field_files) >= 4)
            check(f"표 OOXML 디스크 저장 ({len(table_files)}개)", len(table_files) == 2)
        else:
            bugs.append("raw 디스크 폴더 없음")
            check("raw 디스크 폴더 존재", False)

        # ── RENDER ──
        section("RENDER")
        spec = _load_default_spec()
        try:
            out_bytes = render_docx(outline, spec, user_id=user_id, job_id=job_id)
            print(f"  렌더 완료: {len(out_bytes) / 1024:.1f} KB")
            check("렌더 예외 없음", True)
        except Exception as e:
            bugs.append(f"렌더 예외: {e}")
            check("렌더 예외 없음", False, repr(e))
            return 2

        # ── 재파싱 검증 ──
        section("검증 #5 — 렌더 결과 재오픈")
        try:
            out_doc = Document(io.BytesIO(out_bytes))
            check("렌더 결과 docx 로 재오픈 가능", True)
        except Exception as e:
            bugs.append(f"재오픈 실패: {e}")
            check("렌더 결과 docx 로 재오픈 가능", False, repr(e))
            return 3

        out_tables = out_doc.tables
        if not check("재오픈한 docx 에 표 2개", len(out_tables) == 2, f"실제 {len(out_tables)}"):
            bugs.append(f"렌더 결과 표 개수 미스매치: {len(out_tables)} ≠ 2")

        out_body = out_doc.element.body
        out_fldchars = out_body.findall(f".//{qn('w:fldChar')}")
        out_fldsimples = out_body.findall(f".//{qn('w:fldSimple')}")
        out_bookmarks = out_body.findall(f".//{qn('w:bookmarkStart')}")
        bm_names = [b.get(qn("w:name")) for b in out_bookmarks]
        instr_texts = [(it.text or "") for it in out_body.findall(f".//{qn('w:instrText')}")]

        if not check(
            "재오픈 docx 의 fldChar ≥ 2개 (TOC begin/end)",
            len(out_fldchars) >= 2,
            f"{len(out_fldchars)}",
        ):
            bugs.append(f"fldChar 보존 실패: {len(out_fldchars)}")
        if not check(
            "재오픈 docx 의 fldSimple ≥ 3개 (REF/PAGEREF)",
            len(out_fldsimples) >= 3,
            f"{len(out_fldsimples)}",
        ):
            bugs.append(f"fldSimple 보존 실패: {len(out_fldsimples)}")
        if not check("TOC instr 텍스트 보존", any("TOC" in t for t in instr_texts)):
            bugs.append("TOC instr text 누락")

        for needed in ("_Ref_overview", "_Ref_purpose", "_Ref_visual", "_Ref_quant"):
            if not check(f"북마크 {needed} 보존", needed in bm_names, f"actual={bm_names}"):
                bugs.append(f"북마크 {needed} 누락")

        # ── 검증: 본문 텍스트 일부 라운드트립 ──
        section("검증 #6 — 본문 텍스트 라운드트립")
        out_text = "\n".join(p.text for p in out_doc.paragraphs)
        for needle in ("2026 사내 시스템", "기획부", "외부 발송 측면", "결론"):
            if not check(f"'{needle}' 존재", needle in out_text):
                bugs.append(f"본문 텍스트 '{needle}' 누락")

        # ── 요약 ──
        section("요약")
        if bugs:
            print(f"  ❌ 발견된 버그/이슈: {len(bugs)}개")
            for b in bugs:
                print(f"     - {b}")
            return 1
        print("  ✅ 모든 검증 통과")
        return 0


if __name__ == "__main__":
    sys.exit(main())
