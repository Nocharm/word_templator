"""필드 보존 paragraph reembed 단위 테스트."""

import io
import uuid

from docx import Document
from docx.oxml.ns import qn

from app.parser.parse_docx import parse_docx
from app.renderer.render_docx import render_docx
from tests.fixtures.build_field_sample import build_sample_with_field_and_bookmark


def _spec():
    from tests.test_render import _load_default_spec

    return _load_default_spec()


def test_render_preserves_toc_field(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    src = tmp_path / "f.docx"
    build_sample_with_field_and_bookmark(src)
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(src.read_bytes(), filename="f.docx", user_id=user_id, job_id=job_id)
    out_bytes = render_docx(outline, _spec(), user_id=user_id, job_id=job_id)
    out_doc = Document(io.BytesIO(out_bytes))

    fldchars = out_doc.element.body.findall(f".//{qn('w:fldChar')}")
    assert len(fldchars) >= 2  # at least begin + end


def test_render_preserves_simple_ref_field(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    src = tmp_path / "f.docx"
    build_sample_with_field_and_bookmark(src)
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(src.read_bytes(), filename="f.docx", user_id=user_id, job_id=job_id)
    out_bytes = render_docx(outline, _spec(), user_id=user_id, job_id=job_id)
    out_doc = Document(io.BytesIO(out_bytes))

    fld_simples = out_doc.element.body.findall(f".//{qn('w:fldSimple')}")
    assert len(fld_simples) >= 1
    assert "REF" in (fld_simples[0].get(qn("w:instr")) or "")


def test_render_preserves_bookmarks_on_heading(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    src = tmp_path / "f.docx"
    build_sample_with_field_and_bookmark(src)
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(src.read_bytes(), filename="f.docx", user_id=user_id, job_id=job_id)
    out_bytes = render_docx(outline, _spec(), user_id=user_id, job_id=job_id)
    out_doc = Document(io.BytesIO(out_bytes))

    bm_starts = out_doc.element.body.findall(f".//{qn('w:bookmarkStart')}")
    names = [b.get(qn("w:name")) for b in bm_starts]
    assert "_Ref100001" in names


def test_render_preserves_toc_instr_text(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    src = tmp_path / "f.docx"
    build_sample_with_field_and_bookmark(src)
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(src.read_bytes(), filename="f.docx", user_id=user_id, job_id=job_id)
    out_bytes = render_docx(outline, _spec(), user_id=user_id, job_id=job_id)
    out_doc = Document(io.BytesIO(out_bytes))

    instr_texts = [it.text or "" for it in out_doc.element.body.findall(f".//{qn('w:instrText')}")]
    assert any("TOC" in t for t in instr_texts)


def test_render_preserves_bookmark_id_pairing(tmp_path, monkeypatch):
    """bookmarkStart 와 bookmarkEnd 가 같은 w:id 로 짝지어져 있어야 한다."""
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    src = tmp_path / "f.docx"
    build_sample_with_field_and_bookmark(src)
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(src.read_bytes(), filename="f.docx", user_id=user_id, job_id=job_id)
    out_bytes = render_docx(outline, _spec(), user_id=user_id, job_id=job_id)
    out_doc = Document(io.BytesIO(out_bytes))

    starts = out_doc.element.body.findall(f".//{qn('w:bookmarkStart')}")
    ends = out_doc.element.body.findall(f".//{qn('w:bookmarkEnd')}")
    start_ids = {b.get(qn("w:id")) for b in starts}
    end_ids = {b.get(qn("w:id")) for b in ends}
    assert "1" in start_ids
    assert start_ids.issubset(end_ids)
