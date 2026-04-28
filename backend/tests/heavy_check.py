"""헤비 시나리오 (10 표 + 10 그림 + TOC + 캡션) 라운드트립 점검."""

import io
import json
import os
import sys
import tempfile
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from app.domain.style_spec import StyleSpec  # noqa: E402
from app.parser.parse_docx import parse_docx  # noqa: E402
from app.renderer.render_docx import render_docx  # noqa: E402
from tests.fixtures.build_heavy_sample import build_heavy_sample  # noqa: E402


def _spec() -> StyleSpec:
    seed = ROOT / "app" / "templates_seed" / "report.json"
    return StyleSpec.model_validate(json.loads(seed.read_text(encoding="utf-8"))["spec"])


def section(t: str) -> None:
    print(f"\n=== {t} ===")


def check(label: str, ok: bool, detail: str = "") -> bool:
    icon = "✅" if ok else "❌"
    print(f"  {icon} {label}{(' — ' + detail) if detail else ''}")
    return ok


def main() -> int:
    bugs: list[str] = []

    with tempfile.TemporaryDirectory() as tmp:
        os.environ["DATA_DIR"] = tmp
        src = Path(tmp) / "heavy.docx"
        build_heavy_sample(src)

        section("픽스처")
        print(f"  생성: {src} ({src.stat().st_size / 1024:.1f} KB)")

        user_id = uuid.uuid4()
        job_id = uuid.uuid4()
        outline = parse_docx(
            src.read_bytes(), filename="heavy.docx", user_id=user_id, job_id=job_id
        )
        blocks = outline.blocks

        section("PARSE 통계")
        kinds = {
            k: sum(1 for b in blocks if b.kind == k)
            for k in ("paragraph", "table", "image", "field")
        }
        print(f"  총 블록: {len(blocks)}")
        print(
            f"  kind: paragraph={kinds['paragraph']} table={kinds['table']} image={kinds['image']}"
        )

        levels = {}
        for b in blocks:
            if b.kind == "paragraph":
                levels[b.level] = levels.get(b.level, 0) + 1
        print(f"  level 분포: {dict(sorted(levels.items()))}")

        # ── 표 ──
        section("검증 #1 — 표")
        tables = [b for b in blocks if b.kind == "table"]
        if not check("표 10개 추출", len(tables) == 10, f"실제 {len(tables)}"):
            bugs.append(f"표 개수 미스매치: {len(tables)} ≠ 10")

        with_caption = [t for t in tables if t.caption]
        if not check(
            "표 캡션 흡수 ≥ 9개 (10 중 1개는 헤더만 있어 인접 캡션 없음)",
            len(with_caption) >= 9,
            f"{len(with_caption)}",
        ):
            bugs.append(f"표 캡션 흡수 부족: {len(with_caption)} < 9")

        for i, t in enumerate(tables):
            if not (t.markdown and t.raw_ref):
                bugs.append(f"표[{i}] markdown / raw_ref 비어있음")
        check("표 모두 markdown + raw_ref 채워짐", all(t.markdown and t.raw_ref for t in tables))

        merged = sum(1 for t in tables if t.markdown and "[병합셀]" in t.markdown)
        if not check("병합셀 마커 ≥ 2개", merged >= 2, f"실제 {merged}"):
            bugs.append(f"병합셀 마커 부족: {merged} < 2")

        # 흡수된 캡션 paragraph 잔존 확인
        leftover_table_caps = [
            b
            for b in blocks
            if b.kind == "paragraph"
            and b.text
            and b.text.strip().startswith("표 ")
            and any(t.caption == b.text.strip() for t in tables)
        ]
        if not check(
            "흡수된 표 캡션 paragraph 가 outline 에 잔존하지 않음",
            len(leftover_table_caps) == 0,
            f"잔존 {len(leftover_table_caps)}",
        ):
            bugs.append(f"표 캡션 paragraph 중복: {[b.text for b in leftover_table_caps]}")

        # ── 이미지 ──
        section("검증 #2 — 이미지")
        images = [b for b in blocks if b.kind == "image"]
        if not check("이미지 10개 추출", len(images) == 10, f"실제 {len(images)}"):
            bugs.append(f"이미지 개수 미스매치: {len(images)} ≠ 10")

        with_img_cap = [b for b in images if b.caption]
        if not check("이미지 캡션 흡수 = 10개", len(with_img_cap) == 10, f"{len(with_img_cap)}"):
            bugs.append(f"이미지 캡션 흡수 부족: {len(with_img_cap)} < 10")

        # 모든 이미지에 raw_ref + preview_url
        if not check(
            "이미지 raw_ref + preview_url 모두 설정",
            all(b.raw_ref and b.preview_url for b in images),
        ):
            bugs.append("일부 이미지의 raw_ref/preview_url 누락")

        # 디스크 저장 확인
        img_dir = Path(tmp) / "images" / str(job_id)
        if img_dir.exists():
            saved = list(img_dir.glob("image-*"))
            if not check("이미지 디스크 저장 = 10개", len(saved) == 10, f"{len(saved)}"):
                bugs.append(f"이미지 디스크 파일 부족: {len(saved)} ≠ 10")
        else:
            bugs.append("이미지 디스크 폴더 없음")

        # 흡수된 이미지 캡션 paragraph 잔존 확인
        leftover_fig_caps = [
            b
            for b in blocks
            if b.kind == "paragraph"
            and b.text
            and b.text.strip().startswith("그림 ")
            and any(im.caption == b.text.strip() for im in images)
        ]
        if not check(
            "흡수된 그림 캡션 paragraph 가 outline 에 잔존하지 않음",
            len(leftover_fig_caps) == 0,
            f"잔존 {len(leftover_fig_caps)}",
        ):
            bugs.append(f"그림 캡션 paragraph 중복: {[b.text for b in leftover_fig_caps]}")

        # ── 헤딩 ──
        section("검증 #3 — 헤딩")
        h1 = sum(1 for b in blocks if b.kind == "paragraph" and b.level == 1)
        h2 = sum(1 for b in blocks if b.kind == "paragraph" and b.level == 2)
        if not check("H1 ≥ 5개", h1 >= 5, f"{h1}"):
            bugs.append(f"H1 부족: {h1}")
        if not check("H2 ≥ 4개", h2 >= 4, f"{h2}"):
            bugs.append(f"H2 부족: {h2}")

        cover = next((b for b in blocks if b.text and "표준화 종합 보고서" in b.text), None)
        if not check(
            "표지 제목 휴리스틱 H1",
            bool(cover and cover.level >= 1),
            f"level={cover.level if cover else None}",
        ):
            bugs.append("표지 제목 휴리스틱 실패")

        # ── 필드 ──
        section("검증 #4 — 필드 / 북마크")
        fields = [b for b in blocks if b.kind == "paragraph" and b.raw_xml_ref]
        toc = [b for b in fields if b.field_kind == "toc"]
        check(f"보존 paragraph (필드+북마크) {len(fields)}개", len(fields) >= 5)
        if not check("TOC 1개 분류", len(toc) == 1, f"{len(toc)}"):
            bugs.append(f"TOC 분류 오류: {len(toc)}")

        raw_dir = Path(tmp) / "docs" / str(user_id) / str(job_id) / "raw"
        if raw_dir.exists():
            field_files = list(raw_dir.glob("field-*.xml"))
            table_files = list(raw_dir.glob("table-*.xml"))
            check(f"필드 OOXML 디스크 저장 ({len(field_files)}개)", len(field_files) >= 5)
            if not check(
                "표 OOXML 디스크 저장 = 10개", len(table_files) == 10, f"{len(table_files)}"
            ):
                bugs.append(f"표 OOXML 디스크 파일 부족: {len(table_files)}")
        else:
            bugs.append("raw 디스크 폴더 없음")

        # ── 렌더 ──
        section("RENDER")
        try:
            out_bytes = render_docx(outline, _spec(), user_id=user_id, job_id=job_id)
            print(f"  렌더 완료: {len(out_bytes) / 1024:.1f} KB")
        except Exception as e:
            bugs.append(f"렌더 예외: {e}")
            print(f"  ❌ 렌더 실패: {e!r}")
            return 2

        # ── 재오픈 ──
        section("검증 #5 — 재오픈")
        try:
            out_doc = Document(io.BytesIO(out_bytes))
        except Exception as e:
            bugs.append(f"재오픈 실패: {e}")
            return 3

        if not check("재오픈 docx 표 10개", len(out_doc.tables) == 10, f"{len(out_doc.tables)}"):
            bugs.append(f"렌더 결과 표 개수 미스매치: {len(out_doc.tables)}")

        body = out_doc.element.body
        out_fldchars = body.findall(f".//{qn('w:fldChar')}")
        out_bookmarks = [b.get(qn("w:name")) for b in body.findall(f".//{qn('w:bookmarkStart')}")]
        check(f"TOC fldChar 보존 ({len(out_fldchars)})", len(out_fldchars) >= 2)
        for needed in ("_Ref_overview", "_Ref_plan", "_Ref_data", "_Ref_ref"):
            if not check(f"북마크 {needed} 보존", needed in out_bookmarks):
                bugs.append(f"북마크 {needed} 누락")

        # 본문 텍스트 라운드트립
        text = "\n".join(p.text for p in out_doc.paragraphs)
        for needle in ("표준화 종합 보고서", "정보전략실", "단기 과제", "참고"):
            if not check(f"본문 '{needle}' 존재", needle in text):
                bugs.append(f"본문 '{needle}' 누락")

        # ── 결과 ──
        section("요약")
        if bugs:
            print(f"  ❌ 발견된 버그/이슈: {len(bugs)}개")
            for b in bugs:
                print(f"     - {b}")
            return 1
        print("  ✅ 모든 검증 통과")
        return 0


if __name__ == "__main__":
    sys.exit(main())
