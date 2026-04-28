"""Outline + StyleSpec → .docx round-trip."""

import io
import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.domain.outline import Block, Outline
from app.domain.style_spec import StyleSpec
from app.renderer.render_docx import render_docx

SEED = Path(__file__).resolve().parent.parent / "app" / "templates_seed" / "report.json"


def _load_default_spec() -> StyleSpec:
    raw = json.loads(SEED.read_text(encoding="utf-8"))
    return StyleSpec.model_validate(raw["spec"])


def _build_outline() -> Outline:
    return Outline(
        job_id="j-1",
        source_filename="x.docx",
        blocks=[
            Block(id="b-1", kind="paragraph", level=1, text="개요", detected_by="word_style"),
            Block(
                id="b-2", kind="paragraph", level=0, text="본문 sample.", detected_by="word_style"
            ),
            Block(id="b-3", kind="table", level=0, raw_ref="table-0"),
        ],
    )


def test_render_emits_valid_docx():
    spec = _load_default_spec()
    data = render_docx(_build_outline(), spec)
    doc = Document(io.BytesIO(data))
    paras = [p.text for p in doc.paragraphs]
    assert any("개요" in t for t in paras)
    assert "본문 sample." in paras
    # raw_ref 만 있고 user_id/job_id 없으면 단순 Table 요소로 emit (재파싱 가능 구조 유지)
    assert len(doc.tables) >= 1


def test_render_applies_eastasia_font_for_korean():
    spec = _load_default_spec()
    data = render_docx(_build_outline(), spec)
    doc = Document(io.BytesIO(data))
    # 첫 paragraph "1. 개요" (renumber에 의해 "1. " 붙음)
    p = doc.paragraphs[0]
    run = p.runs[0]
    rPr = run._element.find(qn("w:rPr"))
    rFonts = rPr.find(qn("w:rFonts")) if rPr is not None else None
    assert rFonts is not None
    assert rFonts.get(qn("w:eastAsia")) == "맑은 고딕"
    assert rFonts.get(qn("w:ascii")) == "Arial"


def test_render_respects_block_alignment():
    spec = _load_default_spec()
    outline = Outline(
        job_id="j-1",
        source_filename="x.docx",
        blocks=[
            Block(
                id="b-1",
                kind="paragraph",
                level=0,
                text="가운데",
                detected_by="user",
                alignment="center",
            ),
            Block(id="b-2", kind="paragraph", level=0, text="기본", detected_by="user"),
        ],
    )
    data = render_docx(outline, spec)
    doc = Document(io.BytesIO(data))
    # spec 의 alignment 는 'justify' 인데 b-1 은 'center' 로 override 되어야 함
    p1 = doc.paragraphs[0]
    p2 = doc.paragraphs[1]
    # WD_ALIGN_PARAGRAPH.CENTER == 1, JUSTIFY == 3
    assert int(p1.paragraph_format.alignment) == 1
    assert int(p2.paragraph_format.alignment) == 3
