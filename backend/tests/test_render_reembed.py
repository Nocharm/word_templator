"""표 OOXML 재삽입 + 스타일 오버라이드 단위 테스트."""

import io
import json
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.domain.outline import Block, Outline
from app.domain.style_spec import StyleSpec
from app.parser.parse_docx import parse_docx
from app.renderer.render_docx import render_docx
from tests.fixtures.build_table_image_sample import build_sample_with_table_and_image

_SEED = Path(__file__).resolve().parent.parent / "app" / "templates_seed" / "report.json"


def _spec_minimal() -> StyleSpec:
    raw = json.loads(_SEED.read_text(encoding="utf-8"))
    return StyleSpec.model_validate(raw["spec"])


def test_render_reembeds_table_oxml(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    src = tmp_path / "s.docx"
    build_sample_with_table_and_image(src)
    content = src.read_bytes()
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(content, filename="s.docx", user_id=user_id, job_id=job_id)

    out_bytes = render_docx(outline, _spec_minimal(), user_id=user_id, job_id=job_id)
    out_doc = Document(io.BytesIO(out_bytes))

    tables = out_doc.tables
    assert len(tables) == 1
    cells = [c.text.strip() for row in tables[0].rows for c in row.cells]
    assert "구분" in cells and "값" in cells and "10" in cells


def test_render_applies_table_border_color(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    src = tmp_path / "s.docx"
    build_sample_with_table_and_image(src)
    content = src.read_bytes()
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(content, filename="s.docx", user_id=user_id, job_id=job_id)

    spec = _spec_minimal()
    spec.table.border_color = "#FF0000"
    spec.table.border_width_pt = 1.0

    out_bytes = render_docx(outline, spec, user_id=user_id, job_id=job_id)
    out_doc = Document(io.BytesIO(out_bytes))
    tbl = out_doc.tables[0]
    borders = tbl._tbl.findall(f".//{qn('w:tblBorders')}/{qn('w:top')}")
    assert borders, "tblBorders/w:top not found"
    assert borders[0].get(qn("w:color")) == "FF0000"


def test_render_image_placeholder_when_block_image(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    outline = Outline(
        job_id="x",
        source_filename="x.docx",
        blocks=[
            Block(
                id="b-1",
                kind="image",
                level=0,
                caption="그림 1. 예시",
                raw_ref="image-0",
                preview_url="/api/jobs/none/images/0",
            )
        ],
    )
    out_bytes = render_docx(outline, _spec_minimal())
    out_doc = Document(io.BytesIO(out_bytes))
    paras = [p.text for p in out_doc.paragraphs]
    assert any("[이미지]" in t and "그림 1. 예시" in t for t in paras)


def test_render_reembed_falls_back_on_corrupt_raw_ref(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    # Write a corrupt fragment to the raw_ref path
    from app.storage.files import raw_ooxml_path

    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    raw_ooxml_path(user_id, job_id, "table-0").write_bytes(b"not valid xml")

    outline = Outline(
        job_id=str(job_id),
        source_filename="x.docx",
        blocks=[
            Block(
                id="b-1",
                kind="table",
                level=0,
                raw_ref="table-0",
                markdown="| a |\n| - |\n| b |",
            )
        ],
    )
    out_bytes = render_docx(outline, _spec_minimal(), user_id=user_id, job_id=job_id)
    out_doc = Document(io.BytesIO(out_bytes))
    paras = [p.text for p in out_doc.paragraphs]
    assert any("[표 원본 누락 — table-0]" in t for t in paras)
