"""필드(TOC/REF) + 북마크가 들어있는 .docx 픽스처 빌더."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree


def add_bookmark(paragraph: Paragraph, *, name: str, bm_id: int) -> None:
    """문단 시작/끝에 <w:bookmarkStart/> / <w:bookmarkEnd/> 를 끼워넣는다."""
    p_el = paragraph._p
    start = etree.Element(qn("w:bookmarkStart"))
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), name)
    end = etree.Element(qn("w:bookmarkEnd"))
    end.set(qn("w:id"), str(bm_id))
    pPr = p_el.find(qn("w:pPr"))
    insert_idx = 0 if pPr is None else list(p_el).index(pPr) + 1
    p_el.insert(insert_idx, start)
    p_el.append(end)


def add_simple_field(paragraph: Paragraph, *, instr: str, display_text: str) -> None:
    """<w:fldSimple w:instr="...">display</w:fldSimple> 형태로 삽입."""
    p_el = paragraph._p
    fld = etree.SubElement(p_el, qn("w:fldSimple"))
    fld.set(qn("w:instr"), instr)
    r = etree.SubElement(fld, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = display_text
    t.set(qn("xml:space"), "preserve")


def add_complex_toc_field(paragraph: Paragraph, *, display_text: str) -> None:
    """3-part complex field: begin / instrText / separate / display / end."""
    p_el = paragraph._p
    instr = 'TOC \\o "1-3" \\h \\z \\u'

    def run_with_fldchar(t: str) -> None:
        r = etree.SubElement(p_el, qn("w:r"))
        fc = etree.SubElement(r, qn("w:fldChar"))
        fc.set(qn("w:fldCharType"), t)

    run_with_fldchar("begin")
    r2 = etree.SubElement(p_el, qn("w:r"))
    it = etree.SubElement(r2, qn("w:instrText"))
    it.text = instr
    it.set(qn("xml:space"), "preserve")
    run_with_fldchar("separate")
    r3 = etree.SubElement(p_el, qn("w:r"))
    t = etree.SubElement(r3, qn("w:t"))
    t.text = display_text
    t.set(qn("xml:space"), "preserve")
    run_with_fldchar("end")


def build_sample_with_field_and_bookmark(out_path: Path) -> Path:
    doc = Document()
    h1 = doc.add_paragraph("개요", style="Heading 1")
    add_bookmark(h1, name="_Ref100001", bm_id=1)

    toc_para = doc.add_paragraph()
    add_complex_toc_field(toc_para, display_text="목차")

    body = doc.add_paragraph()
    body.add_run("자세한 내용은 ")
    add_simple_field(body, instr="REF _Ref100001 \\h", display_text="개요")
    body.add_run(" 참조.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    here = Path(__file__).resolve().parent
    print(f"wrote {build_sample_with_field_and_bookmark(here / 'sample_field.docx')}")
