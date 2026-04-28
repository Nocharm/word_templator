"""표 10개 / 이미지 10개 / TOC + 캡션이 들어간 헤비 비정규화 .docx 픽스처.

다양한 시나리오:
- 페이지 목차 (TOC complex field, 다수 엔트리)
- 표 캡션 위/아래 혼합, 일부 병합셀, 일부 빈 셀
- 이미지 캡션 위/아래 혼합, 일부 캡션 두 줄
- 제목 휴리스틱: 1./1.1./(1)/①/한글/Heading 1 혼재
- 폰트 패밀리/사이즈 혼재 (맑은 고딕/Arial/돋움/12/14/16)
- 머리글/꼬리글 텍스트 시뮬레이션 (그냥 본문에 박음)

실행: python tests/fixtures/build_heavy_sample.py
산출: tests/fixtures/sample_heavy.docx
"""

import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from lxml import etree

OUT = Path(__file__).resolve().parent

_TINY_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452"
    "0000001000000010080600000001ff3f"
    "ff0000003c49444154789c63f8ff9f81"
    "8181818181818181c104f8ffff3f0334"
    "1aff8181818181818181c104f8ffffff"
    "ff019f0181818181818181c184820301"
    "00ad1503e1d2e7e6e90000000049454e"
    "44ae426082"
)


def add_run(p, text, *, bold=False, italic=False, size_pt=None, font=None, color=None):
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size_pt is not None:
        r.font.size = Pt(size_pt)
    if font:
        r.font.name = font
    if color:
        r.font.color.rgb = RGBColor(*color)


def add_para(doc, text="", *, bold=False, italic=False, size_pt=None, font=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        add_run(p, text, bold=bold, italic=italic, size_pt=size_pt, font=font)
    return p


def add_bookmark(p: Paragraph, *, name: str, bm_id: int) -> None:
    p_el = p._p
    start = etree.Element(qn("w:bookmarkStart"))
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), name)
    end = etree.Element(qn("w:bookmarkEnd"))
    end.set(qn("w:id"), str(bm_id))
    pPr = p_el.find(qn("w:pPr"))
    insert_idx = 0 if pPr is None else list(p_el).index(pPr) + 1
    p_el.insert(insert_idx, start)
    p_el.append(end)


def add_complex_toc_field(p: Paragraph, *, display: str) -> None:
    p_el = p._p
    instr = 'TOC \\o "1-4" \\h \\z \\u'

    def fc(t: str) -> None:
        run = etree.SubElement(p_el, qn("w:r"))
        ch = etree.SubElement(run, qn("w:fldChar"))
        ch.set(qn("w:fldCharType"), t)

    fc("begin")
    r2 = etree.SubElement(p_el, qn("w:r"))
    it = etree.SubElement(r2, qn("w:instrText"))
    it.text = instr
    it.set(qn("xml:space"), "preserve")
    fc("separate")
    r3 = etree.SubElement(p_el, qn("w:r"))
    t = etree.SubElement(r3, qn("w:t"))
    t.text = display
    t.set(qn("xml:space"), "preserve")
    fc("end")


def add_pagebreak(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = etree.SubElement(r._r, qn("w:br"))
    br.set(qn("w:type"), "page")


def add_table(doc, rows, *, header=True, merge_first_two_cols=False):
    """rows: list of lists. 첫 행을 헤더로 본다."""
    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.cell(ri, ci).text = str(val)
    if merge_first_two_cols and len(rows) > 0 and cols >= 2:
        tbl.cell(0, 0).merge(tbl.cell(0, 1))
    if header:
        for ci in range(cols):
            for run in tbl.cell(0, ci).paragraphs[0].runs:
                run.bold = True
    return tbl


def add_image(doc, *, width_inches=1.5):
    doc.add_picture(io.BytesIO(_TINY_PNG), width=Inches(width_inches))


_BODY = (
    "본 절에서는 사내 시스템 표준화 추진 과정에서 도출된 주요 이슈와 "
    "그에 따른 개선 방향을 기술한다. "
    "특히 다양한 부서의 의견을 종합 정리하여 현황·문제점·개선안을 "
    "단계적으로 제시하고자 한다."
)
_LONG = (
    "또한 향후 자동화 도구를 통해 본 문서를 표준 템플릿으로 변환할 때 표·이미지·캡션·상호 참조가 "
    "원본 그대로 유지되어야 한다는 운영 측면의 제약을 함께 고려한다. 인쇄 및 PDF 변환 시 페이지 "
    "번호와 목차가 자동 갱신되어야 하기 때문이다."
)


def build_heavy_sample(out_path: Path) -> Path:
    doc = Document()

    # ── 표지 ──
    add_para(
        doc,
        "사내 시스템 표준화 종합 보고서",
        bold=True,
        size_pt=24,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(doc, "정보전략실 / 운영지원팀", size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "발행일: 2026.04.27", size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_pagebreak(doc)

    # ── 목차 ──
    add_para(doc, "목차", bold=True, size_pt=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    toc = doc.add_paragraph()
    add_complex_toc_field(
        toc,
        display=(
            "1. 개요 ............................................. 3\n"
            "2. 현황 분석 ....................................... 4\n"
            "  2.1. 부서별 현황 ............................... 5\n"
            "  2.2. 시스템별 현황 ............................. 7\n"
            "3. 개선 방향 ....................................... 9\n"
            "  3.1. 단기 과제 .................................. 10\n"
            "  3.2. 중장기 과제 ............................... 12\n"
            "4. 자료 (표·그림) ................................. 14\n"
            "5. 참고 ........................................... 22\n"
        ),
    )
    add_pagebreak(doc)

    # ── 1. 개요 (정상 Heading 1) ──
    h1_overview = doc.add_heading("1. 개요", level=1)
    add_bookmark(h1_overview, name="_Ref_overview", bm_id=1)
    for _ in range(2):
        add_para(doc, _BODY)
        add_para(doc, _LONG)
    add_pagebreak(doc)

    # ── 2. 현황 분석 (manual H1) ──
    add_para(doc, "2. 현황 분석", bold=True, size_pt=18, font="맑은 고딕")
    add_para(doc, _BODY, font="Arial", size_pt=11)
    add_para(doc, _LONG, font="돋움", size_pt=11)

    add_para(doc, "2.1. 부서별 현황", bold=True, size_pt=14)
    add_para(doc, _BODY)

    # 표 1 — 캡션 위
    add_para(doc, "표 1. 부서별 시스템 사용 현황", bold=True, size_pt=11)
    add_table(
        doc,
        rows=[
            ["부서", "주 사용 시스템", "사용자 수"],
            ["기획부", "ERP / 그룹웨어", "12"],
            ["재무팀", "ERP / 회계", "8"],
            ["개발팀", "Git / Jira / 그룹웨어", "20"],
        ],
    )
    add_para(doc)

    add_para(doc, "(1) ERP 사용 패턴", bold=True, size_pt=12)
    add_para(doc, _BODY)
    add_para(doc, "(2) 그룹웨어 사용 패턴", bold=True, size_pt=12)
    add_para(doc, _LONG)

    # 그림 1 — 캡션 아래
    add_image(doc, width_inches=1.6)
    add_para(doc, "그림 1. 시스템 사용 분포", italic=True, size_pt=10)
    add_para(doc, _BODY)

    # 표 2 — 캡션 아래, 빈 셀 포함
    add_table(
        doc,
        rows=[
            ["항목", "현재", "목표"],
            ["일관성", "낮음", "높음"],
            ["속도", "보통", ""],  # 빈 셀
            ["오류율", "", "낮음"],
        ],
    )
    add_para(doc, "표 2. 핵심 지표 (현재 vs 목표)", italic=True, size_pt=10)
    add_para(doc, _LONG)

    add_pagebreak(doc)

    add_para(doc, "2.2. 시스템별 현황", bold=True, size_pt=14)
    add_para(doc, _BODY)

    # 그림 2 — 캡션 위
    add_para(doc, "그림 2. 시스템 간 데이터 흐름")
    add_image(doc, width_inches=2.0)
    add_para(doc, _BODY)

    # 표 3 — 병합셀 + 캡션 위
    add_para(doc, "표 3. 시스템별 책임 부서", bold=True, size_pt=11)
    add_table(
        doc,
        rows=[
            ["구분", "책임", "비고"],
            ["ERP", "재무팀", "야간 점검"],
            ["Git", "개발팀", "주 1회"],
        ],
        merge_first_two_cols=True,
    )
    add_para(doc)

    # 그림 3 — 캡션 두 줄
    add_image(doc, width_inches=1.5)
    add_para(doc, "그림 3. 운영 환경 다이어그램")
    add_para(doc, "(상단: 운영, 하단: 개발/스테이지)", italic=True, size_pt=9)
    add_para(doc, _LONG)

    add_pagebreak(doc)

    # ── 3. 개선 방향 ──
    h1_plan = doc.add_heading("3. 개선 방향", level=1)
    add_bookmark(h1_plan, name="_Ref_plan", bm_id=2)
    add_para(doc, _BODY)

    add_para(doc, "3.1. 단기 과제", bold=True, size_pt=14)
    add_para(doc, "① 양식 통일", bold=True, size_pt=12)
    add_para(doc, _BODY)
    add_para(doc, "② 자동화 도구 도입", bold=True, size_pt=12)
    add_para(doc, _LONG)

    # 표 4
    add_para(doc, "표 4. 단기 과제 일정", bold=True, size_pt=11)
    add_table(
        doc,
        rows=[
            ["과제", "착수", "완료"],
            ["양식 통일", "5월", "6월"],
            ["자동화 도구", "6월", "8월"],
        ],
    )
    add_para(doc)

    # 그림 4
    add_image(doc, width_inches=1.5)
    add_para(doc, "그림 4. 단기 로드맵")
    add_para(doc, _BODY)

    # 표 5 (캡션 아래)
    add_table(
        doc,
        rows=[
            ["KPI", "기준", "목표"],
            ["재가공 시간", "3h", "1h"],
            ["오류 보고", "12건/월", "3건/월"],
        ],
    )
    add_para(doc, "표 5. 단기 KPI", italic=True, size_pt=10)
    add_para(doc, _LONG)

    add_para(doc, "3.2. 중장기 과제", bold=True, size_pt=14)
    add_para(doc, _BODY)

    # 그림 5 (캡션 위)
    add_para(doc, "그림 5. 중장기 청사진")
    add_image(doc, width_inches=2.2)
    add_para(doc, _LONG)

    # 표 6 (병합 헤더)
    add_para(doc, "표 6. 중장기 마일스톤", bold=True, size_pt=11)
    add_table(
        doc,
        rows=[
            ["분기", "목표 산출물", "검증"],
            ["3Q", "표준 템플릿 v1", "QA"],
            ["4Q", "자동화 v1", "베타"],
        ],
        merge_first_two_cols=True,
    )
    add_para(doc)

    add_pagebreak(doc)

    # ── 4. 자료 (표·그림 다수) ──
    h1_data = doc.add_heading("4. 자료", level=1)
    add_bookmark(h1_data, name="_Ref_data", bm_id=3)
    add_para(doc, _BODY)

    add_para(doc, "*** 부록 A. 추가 표 ***", bold=True, size_pt=14)

    # 표 7
    add_para(doc, "표 7. 부서별 만족도", bold=True, size_pt=11)
    add_table(
        doc,
        rows=[
            ["부서", "만족도", "응답수"],
            ["기획", "4.2/5", "12"],
            ["재무", "3.8/5", "8"],
            ["개발", "4.5/5", "20"],
        ],
    )
    add_para(doc)

    # 표 8 — 빈 표 헤더만
    add_para(doc, "표 8. 향후 측정 항목", bold=True, size_pt=11)
    add_table(doc, rows=[["측정 항목", "주기", "도구"]])
    add_para(doc)

    # 표 9
    add_table(
        doc,
        rows=[
            ["연도", "예산", "집행률"],
            ["2024", "1.2억", "95%"],
            ["2025", "1.5억", "92%"],
            ["2026", "1.8억", "—"],
        ],
    )
    add_para(doc, "표 9. 연도별 예산", italic=True, size_pt=10)
    add_para(doc, _BODY)

    # 표 10 — 큰 표
    add_para(doc, "표 10. 종합 평가표", bold=True, size_pt=11)
    add_table(
        doc,
        rows=[
            ["항목", "전", "후", "변화"],
            ["일관성", "낮음", "높음", "↑"],
            ["속도", "보통", "빠름", "↑"],
            ["오류율", "12%", "3%", "↓"],
            ["만족도", "3.5", "4.3", "↑"],
        ],
    )
    add_para(doc)

    add_para(doc, "*** 부록 B. 추가 그림 ***", bold=True, size_pt=14)

    # 그림 6
    add_image(doc, width_inches=1.6)
    add_para(doc, "그림 6. 변환 흐름도")
    add_para(doc, _BODY)

    # 그림 7 (위)
    add_para(doc, "그림 7. 부서별 사용 빈도")
    add_image(doc, width_inches=1.4)

    # 그림 8 (아래)
    add_image(doc, width_inches=1.4)
    add_para(doc, "그림 8. 표준 템플릿 미리보기", italic=True, size_pt=10)

    # 그림 9
    add_para(doc, "그림 9. 시스템 통합 구조")
    add_image(doc, width_inches=2.0)
    add_para(doc, _LONG)

    # 그림 10
    add_image(doc, width_inches=1.5)
    add_para(doc, "그림 10. 마무리 정리도")
    add_para(doc, _BODY)
    add_para(doc, _LONG)

    add_pagebreak(doc)

    # ── 5. 참고 ──
    h1_ref = doc.add_heading("5. 참고", level=1)
    add_bookmark(h1_ref, name="_Ref_ref", bm_id=4)
    add_para(doc, _BODY)
    add_para(doc, _LONG)
    for _ in range(3):
        add_para(doc, _BODY)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    p = build_heavy_sample(OUT / "sample_heavy.docx")
    size_kb = p.stat().st_size / 1024
    print(f"wrote {p} ({size_kb:.1f} KB)")
