"""20페이지 분량의 비정규화 워드 문서 빌더 — 전체 파이프라인 통합 테스트용.

목적: 한 명의 워드 비숙련자가 작성한 듯한 보고서. 헤딩/본문 스타일 혼재,
표·이미지·캡션·필드·북마크·상호참조까지 한 번에 다룬다.

실행: python tests/fixtures/build_full_messy_sample.py
산출: tests/fixtures/sample_full_messy.docx
"""

import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from lxml import etree

OUT = Path(__file__).resolve().parent

_TINY_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452"
    "0000001000000010080600000001ff3f"
    "ff0000003c49444154789c63f8ff9f81"
    "8181818181818181c104f8ffff3f0334"
    "1aff8181818181818181c104f8ffffff"
    "ff019f0181818181818181c184820301"
    "00ad1503e1d2e7e6e90000000049454e"
    "44ae426082"
)


def add_run(
    p: Paragraph, text: str, *, bold=False, italic=False, size_pt=None, font=None, color=None
) -> None:
    r = p.add_run(text)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if size_pt is not None:
        r.font.size = Pt(size_pt)
    if font:
        r.font.name = font
    if color:
        r.font.color.rgb = RGBColor(*color)


def add_para(doc, text="", *, bold=False, size_pt=None, font=None, align=None) -> Paragraph:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        add_run(p, text, bold=bold, size_pt=size_pt, font=font)
    return p


def add_bookmark(p: Paragraph, *, name: str, bm_id: int) -> None:
    p_el = p._p
    start = etree.Element(qn("w:bookmarkStart"))
    start.set(qn("w:id"), str(bm_id))
    start.set(qn("w:name"), name)
    end = etree.Element(qn("w:bookmarkEnd"))
    end.set(qn("w:id"), str(bm_id))
    pPr = p_el.find(qn("w:pPr"))
    insert_idx = 0 if pPr is None else list(p_el).index(pPr) + 1
    p_el.insert(insert_idx, start)
    p_el.append(end)


def add_simple_field(p: Paragraph, *, instr: str, display: str) -> None:
    fld = etree.SubElement(p._p, qn("w:fldSimple"))
    fld.set(qn("w:instr"), instr)
    r = etree.SubElement(fld, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = display
    t.set(qn("xml:space"), "preserve")


def add_complex_toc_field(p: Paragraph, *, display: str) -> None:
    p_el = p._p
    instr = 'TOC \\o "1-3" \\h \\z \\u'

    def fc(t: str) -> None:
        run = etree.SubElement(p_el, qn("w:r"))
        ch = etree.SubElement(run, qn("w:fldChar"))
        ch.set(qn("w:fldCharType"), t)

    fc("begin")
    r2 = etree.SubElement(p_el, qn("w:r"))
    it = etree.SubElement(r2, qn("w:instrText"))
    it.text = instr
    it.set(qn("xml:space"), "preserve")
    fc("separate")
    r3 = etree.SubElement(p_el, qn("w:r"))
    t = etree.SubElement(r3, qn("w:t"))
    t.text = display
    t.set(qn("xml:space"), "preserve")
    fc("end")


def add_pagebreak(doc) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    br = etree.SubElement(r._r, qn("w:br"))
    br.set(qn("w:type"), "page")


_BODY = (
    "본 문서는 사내 시스템 표준화를 위해 각 팀에서 제출한 의견을 종합하여 정리한 자료입니다. "
    "특히 여러 부서에 걸쳐 사용되는 양식과 워드 문서가 일관되지 않아 외부 발송 시 불필요한 "
    "재가공 작업이 반복되는 점을 해소하고자 함이 1차 목표입니다."
)

_LONG = (
    "또한 향후 자동화 도구를 통해 본 문서를 표준 템플릿으로 변환할 때 표 / 이미지 / 캡션 / "
    "상호 참조가 깨지지 않도록 원본을 그대로 보존해야 한다는 요구사항이 도출되었습니다. "
    "이는 단순한 시각적 통일이 아니라, 인쇄 / PDF 변환 시 페이지 번호와 목차가 자동으로 "
    "갱신되어야 한다는 운영 측면의 제약과도 직결됩니다."
)


def build_full_messy_sample(out_path: Path) -> Path:
    doc = Document()

    # ── 표지 (manual heading: 휴리스틱이 잡아야 함) ──
    add_para(
        doc,
        "2026 사내 시스템 표준화 결과 보고서",
        bold=True,
        size_pt=22,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(doc, "기획부 / 정보기술팀 공동 작성", size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "발행일: 2026년 4월", size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc)
    add_pagebreak(doc)

    # ── 목차 (TOC 필드) ──
    add_para(doc, "목차", bold=True, size_pt=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    toc = doc.add_paragraph()
    add_complex_toc_field(
        toc,
        display=(
            "1. 개요 ......................................... 1\n"
            "2. 추진 배경 .................................... 3\n"
            "3. 주요 내용 .................................... 5\n"
            "4. 기대 효과 .................................... 12\n"
            "5. 결론 ......................................... 18"
        ),
    )
    add_pagebreak(doc)

    # ── 1. 개요 (Word Heading 1 사용 — 정상 케이스) ──
    h1_overview = doc.add_heading("1. 개요", level=1)
    add_bookmark(h1_overview, name="_Ref_overview", bm_id=1)
    for _ in range(2):
        add_para(doc, _BODY)
        add_para(doc, _LONG)

    h2 = doc.add_heading("1.1. 작성 목적", level=2)
    add_bookmark(h2, name="_Ref_purpose", bm_id=2)
    add_para(doc, _BODY)

    add_para(doc, "(1) 가독성 통일", bold=True, size_pt=12)  # 휴리스틱 (1) 패턴
    add_para(doc, _LONG)
    add_para(doc, "(2) 자동화 가능성 확보", bold=True, size_pt=12)
    add_para(doc, _BODY)

    add_pagebreak(doc)

    # ── 2. 추진 배경 (manual H1 - 표지 휴리스틱 외로) ──
    add_para(doc, "2. 추진 배경", bold=True, size_pt=16)
    for _ in range(2):
        add_para(doc, _BODY)
        add_para(doc, _LONG)

    # 한글 한 글자 H2 패턴
    add_para(doc, "가. 외부 발송 측면", bold=True, size_pt=13)
    body_with_ref = doc.add_paragraph()
    add_run(body_with_ref, "외부에 문서를 발송할 때, 앞서 ")
    add_simple_field(body_with_ref, instr="REF _Ref_overview \\h", display="개요")
    add_run(body_with_ref, " 절에서 언급한 가독성 문제로 인해 재가공이 반복됩니다.")

    add_para(doc, "나. 내부 보관 측면", bold=True, size_pt=13)
    add_para(doc, _LONG)

    # 1.1.1 같은 H3 패턴
    add_para(doc, "2.2.1. 워드 사용 패턴 분석", bold=True, size_pt=12)
    add_para(doc, _BODY)
    add_para(doc, "2.2.2. 도구 선정 기준", bold=True, size_pt=12)
    add_para(doc, _LONG)

    add_pagebreak(doc)

    # ── 3. 주요 내용 — 표 + 이미지 + 캡션 ──
    doc.add_heading("3. 주요 내용", level=1)
    add_para(doc, _BODY)

    # 표 (캡션 위)
    add_para(doc, "표 1. 부서별 의견 분류", bold=True, size_pt=11)
    tbl1 = doc.add_table(rows=4, cols=3)
    tbl1.style = "Table Grid"
    for col, hdr in enumerate(["부서", "주요 의견", "건수"]):
        c = tbl1.cell(0, col)
        c.text = hdr
    rows = [
        ("기획부", "양식 통일 필요", "12"),
        ("재무팀", "표 형식 표준화", "8"),
        ("개발팀", "버전 관리 도구 연동", "5"),
    ]
    for ri, row_vals in enumerate(rows, start=1):
        for ci, v in enumerate(row_vals):
            tbl1.cell(ri, ci).text = v

    add_para(doc)
    add_para(doc, "위 표에서 보듯, 양식 / 표 형식 표준화 의견이 전체의 70% 이상을 차지하였습니다.")
    add_para(doc, _LONG)

    h2_3_2 = doc.add_heading("3.2. 시각 자료", level=2)
    add_bookmark(h2_3_2, name="_Ref_visual", bm_id=3)
    add_para(doc, _BODY)

    # 이미지 + 캡션 (아래)
    add_para(doc, "그림 1. 표준화 전 / 후 비교 다이어그램")
    doc.add_picture(io.BytesIO(_TINY_PNG), width=Inches(2.0))
    add_para(doc, "위 그림은 변환 전 / 후 폰트와 헤딩 구조의 차이를 한 눈에 보여줍니다.")

    add_pagebreak(doc)

    # 표 (병합셀 포함, 캡션 아래)
    tbl2 = doc.add_table(rows=3, cols=4)
    tbl2.style = "Table Grid"
    # 헤더 병합
    tbl2.cell(0, 0).text = "구분"
    tbl2.cell(0, 1).merge(tbl2.cell(0, 2)).text = "측정 항목"
    tbl2.cell(0, 3).text = "비고"
    # 데이터
    tbl2.cell(1, 0).text = "전"
    tbl2.cell(1, 1).text = "포맷 일관성"
    tbl2.cell(1, 2).text = "낮음"
    tbl2.cell(1, 3).text = "수동"
    tbl2.cell(2, 0).text = "후"
    tbl2.cell(2, 1).text = "포맷 일관성"
    tbl2.cell(2, 2).text = "높음"
    tbl2.cell(2, 3).text = "자동"
    add_para(doc, "표 2. 변환 전후 비교", bold=False, size_pt=10)

    add_para(doc, _BODY)
    add_para(doc, _LONG)

    add_pagebreak(doc)

    # ── 4. 기대 효과 (혼합 번호 / 들여쓰기 / 본문 ──
    doc.add_heading("4. 기대 효과", level=1)
    add_para(doc, _BODY)

    add_para(doc, "① 시간 단축", bold=True, size_pt=12)  # ① 휴리스틱
    add_para(doc, _LONG)
    add_para(doc, "② 휴먼 에러 감소", bold=True, size_pt=12)
    add_para(doc, _BODY)
    add_para(doc, "③ 외부 발송 시 인상 개선", bold=True, size_pt=12)
    add_para(doc, _LONG)

    h2_4_2 = doc.add_heading("4.2. 정량 효과", level=2)
    add_bookmark(h2_4_2, name="_Ref_quant", bm_id=4)
    add_para(doc, _BODY)

    body_with_ref2 = doc.add_paragraph()
    add_run(body_with_ref2, "이는 ")
    add_simple_field(body_with_ref2, instr="REF _Ref_visual \\h", display="3.2. 시각 자료")
    add_run(body_with_ref2, " 절의 그림 1과 일치하는 결과입니다. 또한 ")
    add_simple_field(body_with_ref2, instr="PAGEREF _Ref_purpose \\h", display="3")
    add_run(body_with_ref2, " 페이지의 목적과도 부합합니다.")

    add_para(doc, _LONG)

    add_pagebreak(doc)

    # ── 5. 결론 + 부록 ──
    doc.add_heading("5. 결론", level=1)
    add_para(doc, _BODY)
    add_para(doc, _LONG)
    add_para(doc, _BODY)

    add_para(doc, "*** 부록 A. 인터뷰 응답자 명단 ***", bold=True, size_pt=14)  # *** *** 휴리스틱
    add_para(doc, "기획부: 김OO, 이OO, 박OO")
    add_para(doc, "재무팀: 정OO, 한OO")
    add_para(doc, "개발팀: 최OO, 윤OO, 강OO, 임OO")
    add_para(doc, _BODY)

    add_para(doc, "*** 부록 B. 약어 정의 ***", bold=True, size_pt=14)
    add_para(doc, "- TOC: Table of Contents (목차)")
    add_para(doc, "- REF: Reference (참조)")
    add_para(doc, "- PAGEREF: Page Reference (페이지 참조)")

    # 더 길게 늘려서 20페이지 근접
    for _ in range(8):
        add_para(doc, _BODY)
        add_para(doc, _LONG)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    p = build_full_messy_sample(OUT / "sample_full_messy.docx")
    size_kb = p.stat().st_size / 1024
    print(f"wrote {p} ({size_kb:.1f} KB)")
