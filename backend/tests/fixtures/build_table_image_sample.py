"""표·이미지·캡션이 들어있는 작은 .docx 픽스처 빌더."""

import io
from pathlib import Path

from docx import Document
from docx.shared import Inches

# 가장 작은 PNG (1x1 RGBA). 필요한 곳에서만 사용.
_TINY_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452"
    "000000010000000108060000001f15c4"
    "890000000b49444154789c63f80f0400"
    "09fb03fdfb5e6b2b0000000049454e44"
    "ae426082"
)


def build_sample_with_table_and_image(out_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("문서 제목", style="Heading 1")

    # 캡션 — 표 위
    doc.add_paragraph("표 1. 결과 요약")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "구분"
    tbl.cell(0, 1).text = "값"
    tbl.cell(1, 0).text = "A"
    tbl.cell(1, 1).text = "10"

    # 본문
    doc.add_paragraph("본문 한 줄.")

    # 이미지 — 캡션은 아래에
    img_io = io.BytesIO(_TINY_PNG)
    doc.add_picture(img_io, width=Inches(1.0))
    doc.add_paragraph("그림 1. 시스템 구조도")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


if __name__ == "__main__":
    here = Path(__file__).resolve().parent
    p = build_sample_with_table_and_image(here / "sample_table_image.docx")
    print(f"wrote {p}")
