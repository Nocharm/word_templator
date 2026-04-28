"""테스트용 .docx fixture 생성. 실행: python tests/fixtures/build_fixtures.py"""

from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(__file__).resolve().parent


def build_simple() -> None:
    """헤딩 + 본문만."""
    doc = Document()
    doc.add_heading("개요", level=1)
    doc.add_paragraph("본 문서는 샘플입니다.")
    doc.add_heading("배경", level=2)
    doc.add_paragraph("Lorem ipsum.")
    doc.save(OUT / "sample_simple.docx")


def build_heuristic() -> None:
    """Word 스타일이 안 적힌, 휴리스틱으로 H1을 잡아야 하는 문서."""
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("1. 개요")
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph("본문.")
    doc.add_paragraph("1.1. 배경")  # H2 휴리스틱
    doc.add_paragraph("내용.")
    doc.save(OUT / "sample_heuristic.docx")


def build_with_table() -> None:
    """표가 섞인 문서 (Phase 1: placeholder block 출력 검증)."""
    doc = Document()
    doc.add_heading("결과", level=1)
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "구분"
    t.cell(0, 1).text = "값"
    t.cell(1, 0).text = "A"
    t.cell(1, 1).text = "10"
    doc.add_paragraph("이상 결과 요약.")
    doc.save(OUT / "sample_with_table.docx")


if __name__ == "__main__":
    build_simple()
    build_heuristic()
    build_with_table()
    print("fixtures generated.")
