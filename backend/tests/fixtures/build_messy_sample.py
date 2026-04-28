"""테스트용 — 워드를 잘 못 쓰는 사람이 작성한 듯한 비표준 .docx 생성.

실행: python tests/fixtures/build_messy_sample.py
산출: tests/fixtures/sample_messy.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent


def add_styled_para(
    doc,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    size_pt: float = 11,
    font_name: str | None = None,
    color_rgb: tuple[int, int, int] | None = None,
    align=None,
):
    """수동으로 폰트/굵기/정렬을 설정한 문단 — Word Heading 스타일은 안 씀."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    r.font.size = Pt(size_pt)
    if font_name:
        r.font.name = font_name
    if color_rgb:
        r.font.color.rgb = RGBColor(*color_rgb)
    return p


def build() -> None:
    doc = Document()

    # ── 1. 표지처럼 보이는 헤딩 (Heading 스타일 ❌, 수동) ──
    add_styled_para(
        doc,
        "2026 사내 워크숍 결과 보고서",
        bold=True,
        size_pt=22,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_styled_para(doc, "기획부 김OO", size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()  # 빈 줄
    doc.add_paragraph()

    # ── 2. 일반 헤딩 + 본문 (제대로 된 Word Heading 1) ──
    doc.add_heading("개요", level=1)
    doc.add_paragraph(
        "본 보고서는 2026년 4월 진행된 사내 워크숍의 결과를 정리하고, " "향후 개선 방향을 제시한다."
    )
    doc.add_paragraph("워크숍은 4월 15일부터 17일까지 강원도 평창에서 진행되었다.")
    doc.add_paragraph()  # 의미 없는 빈 줄

    # ── 3. 가짜 헤딩 — bold + 큰 폰트, Heading 스타일 없음 ──
    add_styled_para(doc, "1. 진행 일정", bold=True, size_pt=15)
    doc.add_paragraph("일정은 다음과 같다.")

    # 들여쓴 글머리 — 손수 만든 리스트
    add_styled_para(doc, "    - 첫째날: 오리엔테이션 및 팀빌딩")
    add_styled_para(doc, "    - 둘째날: 부서별 발표 및 토론")
    add_styled_para(doc, "    - 셋째날: 종합 토의 및 마무리")

    # ── 4. 다양한 번호 형식 ──
    add_styled_para(doc, "1.1 세부 일정", bold=True, size_pt=13)
    add_styled_para(doc, "(1) 09:00 - 09:30: 환영사")
    add_styled_para(doc, "(2) 09:30 - 12:00: 부서 발표")
    add_styled_para(doc, "1) 점심시간: 12:00 - 13:30")
    add_styled_para(doc, "① 오후 1세션: 13:30 - 15:00")
    add_styled_para(doc, "② 오후 2세션: 15:30 - 17:00")

    # ── 5. 한글 마커 (가/나/다 → 휴리스틱 H2) ──
    add_styled_para(doc, "1.2 참가자 구성", bold=True, size_pt=13)
    add_styled_para(doc, "가. 기획부 12명")
    add_styled_para(doc, "나. 개발부 8명")
    add_styled_para(doc, "다. 디자인부 5명")

    # ── 6. 갑자기 폰트가 자꾸 바뀜 ──
    doc.add_heading("주요 발표 내용", level=1)  # 제대로 된 H1

    p = doc.add_paragraph()
    r1 = p.add_run("첫 번째 발표는 ")
    r1.font.name = "맑은 고딕"
    r1.font.size = Pt(11)
    r2 = p.add_run("'클라우드 마이그레이션 사례'")
    r2.font.name = "굴림"
    r2.font.size = Pt(13)
    r2.bold = True
    r3 = p.add_run("였으며, 좋은 반응을 얻었다.")
    r3.font.name = "바탕"
    r3.font.size = Pt(10)

    # 색상이 갑자기 바뀐 강조
    add_styled_para(
        doc,
        "중요: 본 보고서는 사내용이며 외부 공유 금지.",
        bold=True,
        size_pt=12,
        color_rgb=(0xFF, 0x00, 0x00),
    )

    # 한 문단인데 줄바꿈만 함 (\n으로 인위 분리)
    add_styled_para(
        doc,
        "두 번째 발표는 'AI 도입 로드맵'\n"
        "세 번째 발표는 '디자인 시스템 통합'\n"
        "네 번째 발표는 '성과 측정 지표 개편'",
        size_pt=11,
    )

    # ── 7. 표 — 비표준 사용 (캡션이 표 위/아래 다름, 헤더 수동 굵기) ──
    add_styled_para(
        doc,
        "표 1. 부서별 만족도 조사 결과",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size_pt=11,
    )

    t = doc.add_table(rows=5, cols=4)
    headers = ["부서", "응답 인원", "평균 점수", "비고"]
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        cell.text = h
        for cp in cell.paragraphs:
            for cr in cp.runs:
                cr.bold = True
                cr.font.size = Pt(10)

    rows = [
        ("기획부", "12", "4.5", "보통"),
        ("개발부", "8", "4.8", ""),
        ("디자인부", "5", "4.2", "참여율 저조"),
        ("합계", "25", "4.5", ""),
    ]
    for i, row_data in enumerate(rows, start=1):
        for j, val in enumerate(row_data):
            cell = t.cell(i, j)
            cell.text = val

    doc.add_paragraph()

    # ── 8. 가짜 마커 — *** 결론 *** ──
    add_styled_para(
        doc,
        "*** 결론 ***",
        bold=True,
        size_pt=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph("워크숍은 전반적으로 성공적이었으며, 부서 간 협업 강화에 기여하였다.")

    # ── 9. 점 안 찍은 번호 + 화살표 글머리 ──
    add_styled_para(doc, "3 향후 개선 사항", bold=True, size_pt=13)
    add_styled_para(doc, "▶ 일정을 좀 더 여유있게 배정")
    add_styled_para(doc, "▶ 발표 시간 조정 (1인 30분 → 20분)")
    add_styled_para(doc, "▶ 야간 일정 축소")

    # ── 10. 비대칭 들여쓰기와 글머리 점 ──
    add_styled_para(doc, "  · 부서 간 사전 조율 필요", size_pt=11)
    add_styled_para(doc, "    · 발표 자료 표준화", size_pt=11)
    add_styled_para(doc, "      · 슬라이드 템플릿 통일", size_pt=11)

    # ── 11. 표지처럼 마무리 ──
    add_styled_para(
        doc,
        "감사합니다.",
        bold=True,
        size_pt=18,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()
    add_styled_para(doc, "2026. 4. 26.", size_pt=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_styled_para(doc, "기획부 김OO 드림", size_pt=11, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(OUT / "sample_messy.docx")
    print(f"sample_messy.docx generated at {OUT}")


if __name__ == "__main__":
    build()
