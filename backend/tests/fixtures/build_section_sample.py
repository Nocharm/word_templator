"""섹션 분리 + 머리말/꼬리말 fixture.

생성 결과: sample_with_sections.docx
- 섹션 1: portrait, 본문 2 paragraph + 머리말("페이지 머리말") + 꼬리말("Footer L")
- 섹션 break + 페이지 방향 전환 (portrait → landscape)
- 섹션 2: landscape, 본문 2 paragraph

실행: python tests/fixtures/build_section_sample.py
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Mm

OUT = Path(__file__).resolve().parent


def build_sections() -> None:
    doc = Document()

    # 섹션 1: portrait + header/footer
    section1 = doc.sections[0]
    section1.orientation = WD_ORIENT.PORTRAIT
    section1.page_width = Mm(210)
    section1.page_height = Mm(297)
    section1.header.is_linked_to_previous = False
    section1.header.paragraphs[0].text = "페이지 머리말"
    section1.footer.is_linked_to_previous = False
    section1.footer.paragraphs[0].text = "Footer L"

    doc.add_heading("개요", level=1)
    doc.add_paragraph("세로 방향 본문 첫 단락.")

    # 섹션 break + landscape
    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    section2.orientation = WD_ORIENT.LANDSCAPE
    section2.page_width = Mm(297)
    section2.page_height = Mm(210)

    doc.add_heading("부록", level=1)
    doc.add_paragraph("가로 방향 본문 단락.")

    doc.save(OUT / "sample_with_sections.docx")


if __name__ == "__main__":
    build_sections()
    print("section fixture generated.")
