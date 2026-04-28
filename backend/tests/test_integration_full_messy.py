"""20페이지 비정규화 docx 라운드트립 회귀 테스트 — Phase 1~4 통합."""

import io
import json
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.domain.style_spec import StyleSpec
from app.parser.parse_docx import parse_docx
from app.renderer.render_docx import render_docx
from tests.fixtures.build_full_messy_sample import build_full_messy_sample

ROOT = Path(__file__).resolve().parents[1]


def _spec() -> StyleSpec:
    seed = ROOT / "app" / "templates_seed" / "report.json"
    data = json.loads(seed.read_text(encoding="utf-8"))
    return StyleSpec.model_validate(data["spec"])


def _parsed(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    src = tmp_path / "full.docx"
    build_full_messy_sample(src)
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(
        src.read_bytes(),
        filename="full.docx",
        user_id=user_id,
        job_id=job_id,
    )
    return outline, user_id, job_id


def test_full_messy_heading_detection_minimums(tmp_path, monkeypatch):
    outline, _, _ = _parsed(tmp_path, monkeypatch)
    blocks = outline.blocks
    h1 = sum(1 for b in blocks if b.kind == "paragraph" and b.level == 1)
    h2 = sum(1 for b in blocks if b.kind == "paragraph" and b.level == 2)
    h3 = sum(1 for b in blocks if b.kind == "paragraph" and b.level == 3)
    assert h1 >= 5, f"H1 부족: {h1}"
    assert h2 >= 3, f"H2 부족: {h2}"
    assert h3 >= 2, f"H3 부족: {h3}"

    cover = next((b for b in blocks if b.text and "2026 사내 시스템" in b.text), None)
    assert cover is not None
    assert cover.level >= 1, "표지 제목 휴리스틱 실패"

    appendix = [
        b
        for b in blocks
        if b.kind == "paragraph" and b.text and "부록" in b.text and "***" in b.text
    ]
    assert appendix and all(b.level >= 1 for b in appendix), "*** 부록 *** 휴리스틱 실패"


def test_full_messy_table_extraction_and_caption_absorption(tmp_path, monkeypatch):
    outline, _, _ = _parsed(tmp_path, monkeypatch)
    tables = [b for b in outline.blocks if b.kind == "table"]
    assert len(tables) == 2

    cap_above = next((t for t in tables if t.caption and "부서별" in t.caption), None)
    cap_below = next((t for t in tables if t.caption and "변환" in t.caption), None)
    assert cap_above is not None, "캡션 (표 위) 흡수 실패"
    assert cap_below is not None, "캡션 (표 아래) 흡수 실패"

    leftover = [
        b
        for b in outline.blocks
        if b.kind == "paragraph" and b.text and ("표 1. 부서별" in b.text or "표 2. 변환" in b.text)
    ]
    assert leftover == [], "흡수된 캡션이 outline 에 중복으로 남음"

    assert all(t.markdown and t.raw_ref for t in tables)
    assert any(t.markdown and "[병합셀]" in t.markdown for t in tables), "병합셀 마커 누락"


def test_full_messy_image_extraction(tmp_path, monkeypatch):
    outline, _, job_id = _parsed(tmp_path, monkeypatch)
    images = [b for b in outline.blocks if b.kind == "image"]
    assert len(images) == 1
    img = images[0]
    assert img.raw_ref and img.preview_url
    assert img.caption and "그림 1" in img.caption

    img_dir = Path(__import__("os").environ["DATA_DIR"]) / "images" / str(job_id)
    saved = list(img_dir.glob("image-*"))
    assert saved, "이미지 디스크 저장 실패"


def test_full_messy_field_classification(tmp_path, monkeypatch):
    outline, _, _ = _parsed(tmp_path, monkeypatch)
    fields = [b for b in outline.blocks if b.kind == "paragraph" and b.raw_xml_ref]
    toc = [b for b in fields if b.field_kind == "toc"]
    ref = [b for b in fields if b.field_kind == "ref"]
    pageref = [b for b in fields if b.field_kind == "pageref"]
    assert len(fields) >= 4, f"필드/북마크 보존 부족: {len(fields)}"
    assert len(toc) == 1
    # 본문 #1 = REF only, 본문 #2 = REF+PAGEREF mixed (PAGEREF 우선순위로 분류)
    assert len(ref) >= 1
    assert len(pageref) >= 1, "PAGEREF 분류 실패 — 우선순위 이슈"


def test_full_messy_render_roundtrip(tmp_path, monkeypatch):
    outline, user_id, job_id = _parsed(tmp_path, monkeypatch)
    out_bytes = render_docx(outline, _spec(), user_id=user_id, job_id=job_id)
    assert len(out_bytes) > 10_000

    out_doc = Document(io.BytesIO(out_bytes))
    assert len(out_doc.tables) == 2

    body = out_doc.element.body
    assert len(body.findall(f".//{qn('w:fldChar')}")) >= 2, "TOC fldChar 누락"
    assert len(body.findall(f".//{qn('w:fldSimple')}")) >= 3, "REF/PAGEREF fldSimple 누락"

    instr = [(it.text or "") for it in body.findall(f".//{qn('w:instrText')}")]
    assert any("TOC" in t for t in instr), "TOC instr 텍스트 누락"

    bm_names = [b.get(qn("w:name")) for b in body.findall(f".//{qn('w:bookmarkStart')}")]
    for needed in ("_Ref_overview", "_Ref_purpose", "_Ref_visual", "_Ref_quant"):
        assert needed in bm_names, f"북마크 {needed} 누락"

    text = "\n".join(p.text for p in out_doc.paragraphs)
    for needle in ("2026 사내 시스템", "기획부", "외부 발송 측면", "결론"):
        assert needle in text, f"본문 텍스트 '{needle}' 누락"
