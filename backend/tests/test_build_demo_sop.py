"""build_demo_sop — 산출물 검증 + 결정성."""

from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT

from scripts.build_demo_sop import OUTPUT_PATH, build_demo_sop_bytes


@pytest.fixture
def doc_bytes() -> bytes:
    return build_demo_sop_bytes()


@pytest.fixture
def doc(doc_bytes, tmp_path) -> Document:
    p = tmp_path / "demo.docx"
    p.write_bytes(doc_bytes)
    return Document(str(p))


def _heading_levels(doc: Document) -> list[int]:
    levels: list[int] = []
    for p in doc.paragraphs:
        style = (p.style.name or "").strip()
        if style.startswith("Heading "):
            try:
                levels.append(int(style.split()[1]))
            except (IndexError, ValueError):
                continue
    return levels


def test_has_4_levels_of_headings(doc: Document) -> None:
    levels = _heading_levels(doc)
    assert 1 in levels
    assert 2 in levels
    assert 3 in levels
    assert 4 in levels


def test_h1_count_is_5(doc: Document) -> None:
    levels = _heading_levels(doc)
    assert sum(1 for x in levels if x == 1) == 5


def test_table_count_is_5(doc: Document) -> None:
    # 본문 4개 + landscape 섹션 가로 표 1개
    assert len(doc.tables) == 5


def test_image_count_is_3(doc: Document) -> None:
    drawings = doc.element.body.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
    )
    assert len(drawings) == 3


def test_has_landscape_section(doc: Document) -> None:
    orientations = [s.orientation for s in doc.sections]
    assert WD_ORIENT.LANDSCAPE in orientations
    assert WD_ORIENT.PORTRAIT in orientations


def test_has_header_text(doc: Document) -> None:
    for section in doc.sections:
        header_text = "\n".join(p.text for p in section.header.paragraphs)
        if "Demo SOP" in header_text and "시연용" in header_text:
            return
    raise AssertionError("expected 'Demo SOP — Word Templator 시연용' header on at least one section")


def test_deterministic_bytes() -> None:
    a = build_demo_sop_bytes()
    b = build_demo_sop_bytes()
    assert a == b, "build_demo_sop_bytes 가 동일 호출에 동일 바이트를 산출해야 함"


def test_output_path_constant_points_under_templates_seed_demo() -> None:
    assert OUTPUT_PATH.name == "sop_30p.docx"
    assert OUTPUT_PATH.parent.name == "demo"
    assert OUTPUT_PATH.parent.parent.name == "templates_seed"
