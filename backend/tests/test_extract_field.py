"""extract_field 단위 테스트."""

from pathlib import Path

from docx import Document

from app.parser.extract_field import (
    clone_paragraph_xml,
    detect_field_kind,
    extract_bookmarks,
    extract_field_preview,
    paragraph_has_bookmark,
    paragraph_has_field,
)
from tests.fixtures.build_field_sample import build_sample_with_field_and_bookmark


def _load(tmp_path: Path):
    p = tmp_path / "f.docx"
    build_sample_with_field_and_bookmark(p)
    return Document(str(p))


def test_paragraph_has_field_distinguishes_field_paragraphs(tmp_path):
    doc = _load(tmp_path)
    paragraphs = list(doc.paragraphs)
    assert paragraph_has_field(paragraphs[0]) is False
    assert paragraph_has_field(paragraphs[1]) is True
    assert paragraph_has_field(paragraphs[2]) is True


def test_paragraph_has_bookmark_finds_bookmark_on_heading(tmp_path):
    doc = _load(tmp_path)
    paragraphs = list(doc.paragraphs)
    assert paragraph_has_bookmark(paragraphs[0]) is True
    assert paragraph_has_bookmark(paragraphs[1]) is False
    assert paragraph_has_bookmark(paragraphs[2]) is False


def test_detect_field_kind_identifies_toc_and_ref(tmp_path):
    doc = _load(tmp_path)
    paragraphs = list(doc.paragraphs)
    assert detect_field_kind(paragraphs[1]) == "toc"
    assert detect_field_kind(paragraphs[2]) == "ref"
    assert detect_field_kind(paragraphs[0]) is None


def test_detect_field_kind_priority_pageref_over_ref_when_mixed(tmp_path):
    """REF + PAGEREF 가 같은 문단에 섞여 있으면 PAGEREF 가 이긴다 (가장 구체적)."""
    from docx import Document

    from tests.fixtures.build_field_sample import add_simple_field

    p = tmp_path / "mixed.docx"
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("앞 ")
    add_simple_field(para, instr="REF _abc \\h", display_text="개요")
    para.add_run(" 그리고 ")
    add_simple_field(para, instr="PAGEREF _abc \\h", display_text="3")
    doc.save(str(p))

    reopened = list(Document(str(p)).paragraphs)
    assert detect_field_kind(reopened[0]) == "pageref"


def test_extract_field_preview_returns_display_text(tmp_path):
    doc = _load(tmp_path)
    paragraphs = list(doc.paragraphs)
    assert "목차" in (extract_field_preview(paragraphs[1]) or "")
    preview = extract_field_preview(paragraphs[2]) or ""
    assert "개요" in preview


def test_extract_bookmarks_returns_name_and_id(tmp_path):
    doc = _load(tmp_path)
    paragraphs = list(doc.paragraphs)
    bms = extract_bookmarks(paragraphs[0])
    assert len(bms) == 1
    assert bms[0]["name"] == "_Ref100001"
    assert bms[0]["id"] == "1"
    assert extract_bookmarks(paragraphs[2]) == []


def test_clone_paragraph_xml_returns_self_contained_bytes(tmp_path):
    doc = _load(tmp_path)
    paragraphs = list(doc.paragraphs)
    xml = clone_paragraph_xml(paragraphs[1])
    assert isinstance(xml, bytes)
    assert b"<w:p" in xml and b"</w:p>" in xml
    assert b"fldChar" in xml or b"instrText" in xml
