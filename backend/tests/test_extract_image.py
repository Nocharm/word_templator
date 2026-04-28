"""extract_image.iter_image_blobs 단위 테스트."""

from docx import Document
from docx.text.paragraph import Paragraph

from app.parser.extract_image import iter_image_blobs
from tests.fixtures.build_table_image_sample import build_sample_with_table_and_image


def test_iter_image_blobs_returns_png_bytes(tmp_path):
    p = tmp_path / "s.docx"
    build_sample_with_table_and_image(p)
    doc = Document(str(p))

    paragraphs_with_images: list = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            para = Paragraph(child, doc)
            blobs = list(iter_image_blobs(para, doc))
            if blobs:
                paragraphs_with_images.extend(blobs)

    assert len(paragraphs_with_images) == 1
    blob = paragraphs_with_images[0]
    assert blob.content_type.startswith("image/")
    assert blob.ext.lower() in {"png", "jpeg", "jpg"}
    assert blob.data.startswith(b"\x89PNG") or blob.data[:3] == b"\xff\xd8\xff"


def test_iter_image_blobs_empty_paragraph(tmp_path):
    p = tmp_path / "empty.docx"
    doc = Document()
    doc.add_paragraph("그냥 텍스트")
    doc.save(str(p))

    doc2 = Document(str(p))
    para = next(iter(doc2.paragraphs))
    assert list(iter_image_blobs(para, doc2)) == []
