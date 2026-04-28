"""extract_table.table_to_markdown / extract_table.clone_table_xml 단위 테스트."""

from pathlib import Path

from docx import Document

from app.parser.extract_table import clone_table_xml, table_to_markdown
from tests.fixtures.build_table_image_sample import build_sample_with_table_and_image


def _load_first_table(tmp_path: Path):
    p = tmp_path / "s.docx"
    build_sample_with_table_and_image(p)
    doc = Document(str(p))
    return doc.tables[0]


def test_table_to_markdown_renders_header_and_rows(tmp_path):
    table = _load_first_table(tmp_path)
    md = table_to_markdown(table)
    lines = md.splitlines()
    assert lines[0] == "| 구분 | 값 |"
    assert lines[1] == "| --- | --- |"
    assert lines[2] == "| A | 10 |"


def test_clone_table_xml_returns_bytes(tmp_path):
    table = _load_first_table(tmp_path)
    xml = clone_table_xml(table)
    assert isinstance(xml, bytes)
    assert b"<w:tbl" in xml
    assert b"</w:tbl>" in xml


def test_table_to_markdown_handles_merged_cell_marker(tmp_path):
    """gridSpan=2 인 셀은 [병합셀] 표시 — 원본은 깨지지 않게 보존만."""
    p = tmp_path / "merged.docx"
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).merge(tbl.cell(0, 1))
    tbl.cell(0, 0).text = "헤더"
    tbl.cell(1, 0).text = "A"
    tbl.cell(1, 1).text = "B"
    doc.save(str(p))

    md = table_to_markdown(Document(str(p)).tables[0])
    assert "[병합셀]" in md
