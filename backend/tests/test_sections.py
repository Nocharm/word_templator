"""섹션 (page break + orientation) + 머리말/꼬리말 보존 테스트."""

import io
import json
import uuid
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT

from app.domain.style_spec import StyleSpec
from app.parser.parse_docx import parse_docx
from app.renderer.render_docx import render_docx

FIXTURES = Path(__file__).resolve().parent / "fixtures"
SEED = Path(__file__).resolve().parent.parent / "app" / "templates_seed" / "report.json"


def _load_default_spec() -> StyleSpec:
    raw = json.loads(SEED.read_text(encoding="utf-8"))
    return StyleSpec.model_validate(raw["spec"])


def test_parser_detects_two_sections_with_orientations():
    """두 섹션 (portrait + landscape) 을 정확히 분리하고 orientation 보존."""
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(
        (FIXTURES / "sample_with_sections.docx").read_bytes(),
        filename="sample_with_sections.docx",
        user_id=user_id,
        job_id=job_id,
    )

    assert len(outline.sections) == 2, f"expected 2 sections, got {len(outline.sections)}"
    assert outline.sections[0].orientation == "portrait"
    assert outline.sections[1].orientation == "landscape"

    # 섹션 1 에는 "개요" + "세로 방향 본문 첫 단락." 이 들어가야 함
    sec1_block_ids = set(outline.sections[0].block_ids)
    sec1_blocks = [b for b in outline.blocks if b.id in sec1_block_ids]
    sec1_texts = [b.text for b in sec1_blocks if b.kind == "paragraph"]
    assert any("개요" in (t or "") for t in sec1_texts)
    assert any("세로" in (t or "") for t in sec1_texts)

    # 섹션 2 에는 "부록" + "가로 방향 본문 단락." 이 들어가야 함
    sec2_block_ids = set(outline.sections[1].block_ids)
    sec2_blocks = [b for b in outline.blocks if b.id in sec2_block_ids]
    sec2_texts = [b.text for b in sec2_blocks if b.kind == "paragraph"]
    assert any("부록" in (t or "") for t in sec2_texts)
    assert any("가로" in (t or "") for t in sec2_texts)


def test_parser_extracts_header_footer_refs():
    """섹션 1 에 머리말/꼬리말 ref 가 채워지고 디스크에 XML 이 저장된다."""
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(
        (FIXTURES / "sample_with_sections.docx").read_bytes(),
        filename="sample_with_sections.docx",
        user_id=user_id,
        job_id=job_id,
    )

    sec1 = outline.sections[0]
    assert sec1.header_default_ref is not None, "header_default_ref expected on section 1"
    assert sec1.footer_default_ref is not None, "footer_default_ref expected on section 1"

    from app.storage.files import section_part_path

    header_path = section_part_path(user_id, job_id, 0, "header", "default")
    footer_path = section_part_path(user_id, job_id, 0, "footer", "default")
    assert header_path.exists(), f"expected header XML at {header_path}"
    assert footer_path.exists(), f"expected footer XML at {footer_path}"

    header_xml = header_path.read_text(encoding="utf-8")
    assert "페이지 머리말" in header_xml
    footer_xml = footer_path.read_text(encoding="utf-8")
    assert "Footer L" in footer_xml


def test_renderer_preserves_orientation_per_section():
    """렌더 결과가 섹션별로 portrait/landscape 를 유지한다."""
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(
        (FIXTURES / "sample_with_sections.docx").read_bytes(),
        filename="sample_with_sections.docx",
        user_id=user_id,
        job_id=job_id,
    )
    spec = _load_default_spec()
    data = render_docx(outline, spec, user_id=user_id, job_id=job_id)

    out_doc = Document(io.BytesIO(data))
    assert len(out_doc.sections) == 2
    assert out_doc.sections[0].orientation == WD_ORIENT.PORTRAIT
    assert out_doc.sections[1].orientation == WD_ORIENT.LANDSCAPE
    # landscape 섹션은 width > height
    assert out_doc.sections[1].page_width > out_doc.sections[1].page_height


def test_renderer_preserves_header_footer_text():
    """렌더 결과에 원본 머리말/꼬리말 텍스트가 남아있다."""
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(
        (FIXTURES / "sample_with_sections.docx").read_bytes(),
        filename="sample_with_sections.docx",
        user_id=user_id,
        job_id=job_id,
    )
    spec = _load_default_spec()
    data = render_docx(outline, spec, user_id=user_id, job_id=job_id)

    out_doc = Document(io.BytesIO(data))
    sec1 = out_doc.sections[0]
    header_text = "\n".join(p.text for p in sec1.header.paragraphs)
    footer_text = "\n".join(p.text for p in sec1.footer.paragraphs)
    assert "페이지 머리말" in header_text
    assert "Footer L" in footer_text


def test_legacy_outline_without_sections_still_renders():
    """sections 가 비어있는 (legacy) Outline 도 단일 섹션으로 정상 렌더."""
    from app.domain.outline import Block, Outline

    outline = Outline(
        job_id="j-legacy",
        source_filename="legacy.docx",
        blocks=[
            Block(id="b-1", kind="paragraph", level=0, text="legacy body"),
        ],
        # sections 빈 채로
    )
    spec = _load_default_spec()
    data = render_docx(outline, spec)
    out_doc = Document(io.BytesIO(data))
    assert len(out_doc.sections) == 1
    paras = [p.text for p in out_doc.paragraphs]
    assert "legacy body" in paras


def test_parser_simple_doc_yields_single_section():
    """섹션 break 없는 단순 문서는 sections 길이 1 (마지막 body sectPr 1개)."""
    outline = parse_docx(
        (FIXTURES / "sample_simple.docx").read_bytes(),
        filename="sample_simple.docx",
    )
    assert len(outline.sections) == 1
    assert outline.sections[0].orientation == "portrait"
    # 모든 block id 가 단일 섹션에 포함
    section_block_ids = set(outline.sections[0].block_ids)
    all_ids = {b.id for b in outline.blocks}
    assert section_block_ids == all_ids


def test_pgsz_to_mm_conversion():
    """A4 portrait pgSz 가 (210, 297) mm 로 변환된다."""
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(
        (FIXTURES / "sample_with_sections.docx").read_bytes(),
        filename="sample_with_sections.docx",
        user_id=user_id,
        job_id=job_id,
    )
    sec1 = outline.sections[0]
    # python-docx 가 Mm(210) -> 11906 twips ~= 209.99mm 식이라 약간의 반올림 허용.
    assert abs(sec1.page_width_mm - 210) < 1.0
    assert abs(sec1.page_height_mm - 297) < 1.0
    # 섹션 2 (landscape) — sectPr 의 pgSz 는 width=297 이어야 함.
    sec2 = outline.sections[1]
    assert sec2.page_width_mm > sec2.page_height_mm
