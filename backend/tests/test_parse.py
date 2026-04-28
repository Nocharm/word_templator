"""parse_docx: .docx → Outline (Phase 3 표·이미지 + Phase 4 필드/북마크)."""

import uuid
from pathlib import Path

from docx import Document

from app.parser.parse_docx import parse_docx
from tests.fixtures.build_field_sample import build_sample_with_field_and_bookmark
from tests.fixtures.build_table_image_sample import build_sample_with_table_and_image

FIXTURES = Path(__file__).resolve().parent / "fixtures"


def _blank_png_path() -> str:
    """1×1 투명 PNG. 테스트 fixture 에 이미 있으면 그 경로 반환."""
    p = Path(__file__).parent / "fixtures" / "blank.png"
    if not p.exists():
        # 1×1 흰색 PNG (RGB, 8bit)
        p.write_bytes(
            bytes.fromhex(
                "89504e470d0a1a0a0000000d4948445200000001000000010802000000907753de"
                "0000000c49444154789c63f8ffff3f0005fe02fe0def46b80000000049454e44"
                "ae426082"
            )
        )
    return str(p)


def test_parse_simple_headings_and_body():
    outline = parse_docx(
        (FIXTURES / "sample_simple.docx").read_bytes(), filename="sample_simple.docx"
    )
    kinds = [b.kind for b in outline.blocks]
    levels = [b.level for b in outline.blocks if b.kind == "paragraph"]
    assert kinds.count("paragraph") == 4
    # H1, body, H2, body
    assert levels == [1, 0, 2, 0]
    assert outline.blocks[0].detected_by == "word_style"


def test_parse_heuristic_headings():
    outline = parse_docx(
        (FIXTURES / "sample_heuristic.docx").read_bytes(), filename="sample_heuristic.docx"
    )
    paras = [b for b in outline.blocks if b.kind == "paragraph"]
    assert paras[0].level == 1 and paras[0].detected_by == "heuristic"
    assert paras[2].level == 2 and paras[2].detected_by == "heuristic"


def test_parse_table_without_ids_emits_markdown_only():
    outline = parse_docx(
        (FIXTURES / "sample_with_table.docx").read_bytes(), filename="sample_with_table.docx"
    )
    table_blocks = [b for b in outline.blocks if b.kind == "table"]
    assert len(table_blocks) == 1
    # Phase 3: markdown 은 항상 생성, raw_ref 는 user_id/job_id 미지정 시 None.
    assert table_blocks[0].markdown is not None
    assert table_blocks[0].raw_ref is None


def test_parse_block_ids_are_unique():
    outline = parse_docx((FIXTURES / "sample_simple.docx").read_bytes(), filename="x.docx")
    ids = [b.id for b in outline.blocks]
    assert len(ids) == len(set(ids))


def test_parse_collapses_consecutive_empty_paragraphs():
    outline = parse_docx((FIXTURES / "sample_messy.docx").read_bytes(), filename="m.docx")
    paras = [b for b in outline.blocks if b.kind == "paragraph"]
    # consecutive empty 검사
    for i in range(len(paras) - 1):
        a_empty = not (paras[i].text or "").strip()
        b_empty = not (paras[i + 1].text or "").strip()
        assert not (a_empty and b_empty), f"consecutive empty at {i}"


def test_parse_extracts_alignment_for_centered_cover():
    outline = parse_docx((FIXTURES / "sample_messy.docx").read_bytes(), filename="m.docx")
    # 첫 paragraph는 가운데정렬 표지여야 함
    paras = [b for b in outline.blocks if b.kind == "paragraph"]
    assert paras[0].alignment == "center"


# Phase 3 — 표/이미지/캡션 통합 검증.


def test_parse_docx_phase3_extracts_table_image_and_captions(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    p = tmp_path / "s.docx"
    build_sample_with_table_and_image(p)
    content = p.read_bytes()

    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(
        content,
        filename="s.docx",
        user_id=user_id,
        job_id=job_id,
    )

    kinds = [b.kind for b in outline.blocks]
    assert kinds.count("table") == 1
    assert kinds.count("image") == 1

    table = next(b for b in outline.blocks if b.kind == "table")
    assert table.markdown is not None and table.markdown.startswith("| 구분 | 값 |")
    assert table.raw_ref == "table-0"
    assert table.caption == "표 1. 결과 요약"
    assert not any(b.kind == "paragraph" and b.text == "표 1. 결과 요약" for b in outline.blocks)

    image = next(b for b in outline.blocks if b.kind == "image")
    assert image.raw_ref == "image-0"
    assert image.caption == "그림 1. 시스템 구조도"
    assert image.preview_url == f"/api/jobs/{job_id}/images/0"
    assert not any(
        b.kind == "paragraph" and b.text == "그림 1. 시스템 구조도" for b in outline.blocks
    )

    raw_dir = tmp_path / "docs" / str(user_id) / str(job_id) / "raw"
    assert (raw_dir / "table-0.xml").exists()
    img_dir = tmp_path / "images" / str(job_id)
    files = list(img_dir.glob("image-0.*"))
    assert len(files) == 1


def test_parse_docx_phase3_back_compat_without_ids(tmp_path):
    """user_id/job_id 미전달 시에도 파싱은 되며, raw_ref/preview_url 은 None."""
    p = tmp_path / "s.docx"
    build_sample_with_table_and_image(p)
    content = p.read_bytes()
    outline = parse_docx(content, filename="s.docx")
    table = next(b for b in outline.blocks if b.kind == "table")
    assert table.markdown is not None
    assert table.raw_ref is None
    image = next(b for b in outline.blocks if b.kind == "image")
    assert image.raw_ref is None
    assert image.preview_url is None


def test_parse_docx_phase4_preserves_field_and_bookmark_paragraphs(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    src = tmp_path / "f.docx"
    build_sample_with_field_and_bookmark(src)
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(src.read_bytes(), filename="f.docx", user_id=user_id, job_id=job_id)

    paragraphs = [b for b in outline.blocks if b.kind == "paragraph"]
    h1 = next(b for b in paragraphs if b.text and b.text.strip() == "개요")
    assert h1.raw_xml_ref is not None and h1.raw_xml_ref.startswith("field-")
    assert h1.field_kind is None  # bookmark only, no field

    toc = next(b for b in paragraphs if b.field_kind == "toc")
    assert toc.raw_xml_ref is not None
    assert toc.preview_text and "목차" in toc.preview_text

    ref = next(b for b in paragraphs if b.field_kind == "ref")
    assert ref.raw_xml_ref is not None
    assert ref.preview_text and "개요" in ref.preview_text

    raw_dir = tmp_path / "docs" / str(user_id) / str(job_id) / "raw"
    saved = sorted(raw_dir.glob("field-*.xml"))
    assert len(saved) == 3


def test_parse_docx_phase4_no_field_no_save(tmp_path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    p = tmp_path / "plain.docx"
    doc = Document()
    doc.add_paragraph("그냥 텍스트")
    doc.save(str(p))

    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(p.read_bytes(), filename="plain.docx", user_id=user_id, job_id=job_id)
    for b in outline.blocks:
        if b.kind == "paragraph":
            assert b.raw_xml_ref is None
            assert b.field_kind is None


def test_parse_pipeline_marks_heading_skip_and_assigns_captions():
    """파이프라인 통합: heading skip warning + 캡션 자동 번호 + 본문 ref 부착."""
    import io

    from docx import Document

    from app.parser.parse_docx import parse_docx

    doc = Document()
    doc.add_paragraph("큰 제목", style="Heading 1")
    doc.add_paragraph("아주 깊은 제목", style="Heading 3")  # H1 → H3 skip
    doc.add_paragraph("아래 그림 1 을 참조한다.")
    doc.add_picture(_blank_png_path())  # img-1
    doc.add_table(rows=2, cols=2)  # tbl-1, no caption

    buf = io.BytesIO()
    doc.save(buf)
    outline = parse_docx(buf.getvalue(), filename="t.docx")

    headings = [b for b in outline.blocks if b.kind == "paragraph" and b.level >= 1]
    assert any(b.warning == "heading_skip" for b in headings if b.level == 3)

    image_blocks = [b for b in outline.blocks if b.kind == "image"]
    table_blocks = [b for b in outline.blocks if b.kind == "table"]
    img_cap = image_blocks[0].caption
    tbl_cap = table_blocks[0].caption
    assert img_cap is not None and img_cap.startswith("그림 1")
    assert tbl_cap is not None and tbl_cap.startswith("표 1")

    paragraph_with_ref = next(
        b for b in outline.blocks if b.kind == "paragraph" and "그림 1" in (b.text or "")
    )
    assert any(
        r.label_kind == "figure" and r.target_block_id == image_blocks[0].id
        for r in paragraph_with_ref.caption_refs
    )
