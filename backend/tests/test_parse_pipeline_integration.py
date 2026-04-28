"""end-to-end: docx 입력 → outline → docx 출력 → 재파싱 일관성."""

import io

from docx import Document
from lxml import etree

from app.parser.parse_docx import parse_docx
from app.renderer.render_docx import render_docx


def _make_input_docx() -> bytes:
    """캡션 + 본문 참조를 모두 포함하는 docx 생성.

    구조: 제목들 → 본문(표 1 참조) → 표(캡션 "표 1.") 순서로 배치.
    이렇게 해야 본문 단락이 caption_refs 를 갖고, 캡션은 표에 붙는다.
    """
    doc = Document()
    doc.add_paragraph("큰 제목", style="Heading 1")
    doc.add_paragraph("작은 제목", style="Heading 2")
    # 본문에서 표 참조 — 이 단락은 표보다 앞에 있어 캡션으로 흡수되지 않음
    doc.add_paragraph("아래 표 1 을 참조한다.")
    # 표 캡션 (표 바로 앞)
    doc.add_paragraph("표 1. 예시 데이터")
    doc.add_table(rows=2, cols=2)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_caption_emits_seq_field_and_body_emits_ref_field():
    outline = parse_docx(_make_input_docx(), filename="t.docx")
    out_bytes = render_docx_default(outline)

    out_doc = Document(io.BytesIO(out_bytes))
    body_xml = etree.tostring(out_doc.element.body).decode("utf-8")

    assert "SEQ Table" in body_xml or "SEQ Figure" in body_xml
    assert "REF _Ref_" in body_xml
    assert "bookmarkStart" in body_xml


def test_roundtrip_parse_yields_same_caption_numbers():
    input_bytes = _make_input_docx()
    outline_a = parse_docx(input_bytes, filename="t.docx")
    out_bytes = render_docx_default(outline_a)
    outline_b = parse_docx(out_bytes, filename="roundtrip.docx")

    captions_a = [b.caption for b in outline_a.blocks if b.kind in ("table", "image")]
    captions_b = [b.caption for b in outline_b.blocks if b.kind in ("table", "image")]
    assert captions_a == captions_b


def render_docx_default(outline) -> bytes:
    """Wrap render_docx with a default StyleSpec — must mirror parse_docx._default_style_spec."""
    from app.parser.parse_docx import _default_style_spec

    return render_docx(outline, spec=_default_style_spec())
