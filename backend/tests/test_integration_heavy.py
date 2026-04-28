"""헤비 시나리오 (10 표 + 10 이미지 + TOC + 캡션) 회귀 테스트."""

import io
import json
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.domain.style_spec import StyleSpec
from app.parser.parse_docx import parse_docx
from app.renderer.render_docx import render_docx
from tests.fixtures.build_heavy_sample import build_heavy_sample

ROOT = Path(__file__).resolve().parents[1]


def _spec() -> StyleSpec:
    seed = ROOT / "app" / "templates_seed" / "report.json"
    return StyleSpec.model_validate(json.loads(seed.read_text(encoding="utf-8"))["spec"])


def _parsed(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("DATA_DIR", str(tmp_path))
    src = tmp_path / "heavy.docx"
    build_heavy_sample(src)
    user_id = uuid.uuid4()
    job_id = uuid.uuid4()
    outline = parse_docx(src.read_bytes(), filename="heavy.docx", user_id=user_id, job_id=job_id)
    return outline, user_id, job_id


def test_heavy_extracts_ten_tables_with_captions(tmp_path, monkeypatch):
    outline, _, _ = _parsed(tmp_path, monkeypatch)
    tables = [b for b in outline.blocks if b.kind == "table"]
    assert len(tables) == 10, f"표 {len(tables)} ≠ 10"
    captioned = [t for t in tables if t.caption]
    assert len(captioned) >= 9, f"캡션 흡수 {len(captioned)} < 9"
    assert all(t.markdown and t.raw_ref for t in tables)

    merged_markers = sum(1 for t in tables if t.markdown and "[병합셀]" in t.markdown)
    assert merged_markers >= 2, "병합셀 마커 부족"

    leftover = [
        b
        for b in outline.blocks
        if b.kind == "paragraph"
        and b.text
        and b.text.strip().startswith("표 ")
        and any(t.caption == b.text.strip() for t in tables)
    ]
    assert leftover == [], f"흡수된 표 캡션 paragraph 잔존: {[b.text for b in leftover]}"


def test_heavy_extracts_ten_images_with_captions(tmp_path, monkeypatch):
    outline, _, job_id = _parsed(tmp_path, monkeypatch)
    images = [b for b in outline.blocks if b.kind == "image"]
    assert len(images) == 10, f"이미지 {len(images)} ≠ 10"
    captioned = [b for b in images if b.caption]
    assert len(captioned) == 10, f"이미지 캡션 흡수 {len(captioned)} < 10"
    assert all(b.raw_ref and b.preview_url for b in images)

    import os

    img_dir = Path(os.environ["DATA_DIR"]) / "images" / str(job_id)
    saved = list(img_dir.glob("image-*"))
    assert len(saved) == 10, f"이미지 디스크 {len(saved)} ≠ 10"

    leftover = [
        b
        for b in outline.blocks
        if b.kind == "paragraph"
        and b.text
        and b.text.strip().startswith("그림 ")
        and any(im.caption == b.text.strip() for im in images)
    ]
    assert leftover == [], f"흡수된 그림 캡션 paragraph 잔존: {[b.text for b in leftover]}"


def test_heavy_heading_detection(tmp_path, monkeypatch):
    outline, _, _ = _parsed(tmp_path, monkeypatch)
    blocks = outline.blocks
    h1 = sum(1 for b in blocks if b.kind == "paragraph" and b.level == 1)
    h2 = sum(1 for b in blocks if b.kind == "paragraph" and b.level == 2)
    assert h1 >= 5, f"H1 {h1} < 5"
    assert h2 >= 4, f"H2 {h2} < 4"
    cover = next((b for b in blocks if b.text and "표준화 종합 보고서" in b.text), None)
    assert cover and cover.level >= 1, "표지 제목 휴리스틱 실패"


def test_heavy_field_classification(tmp_path, monkeypatch):
    outline, user_id, job_id = _parsed(tmp_path, monkeypatch)
    fields = [b for b in outline.blocks if b.kind == "paragraph" and b.raw_xml_ref]
    toc = [b for b in fields if b.field_kind == "toc"]
    assert len(toc) == 1
    assert len(fields) >= 5, f"필드 보존 {len(fields)} < 5"

    import os

    raw_dir = Path(os.environ["DATA_DIR"]) / "docs" / str(user_id) / str(job_id) / "raw"
    table_files = list(raw_dir.glob("table-*.xml"))
    assert len(table_files) == 10


def test_heavy_render_roundtrip(tmp_path, monkeypatch):
    outline, user_id, job_id = _parsed(tmp_path, monkeypatch)
    out_bytes = render_docx(outline, _spec(), user_id=user_id, job_id=job_id)
    assert len(out_bytes) > 10_000

    out_doc = Document(io.BytesIO(out_bytes))
    assert len(out_doc.tables) == 10, f"렌더 결과 표 {len(out_doc.tables)} ≠ 10"

    body = out_doc.element.body
    assert len(body.findall(f".//{qn('w:fldChar')}")) >= 2
    bm_names = [b.get(qn("w:name")) for b in body.findall(f".//{qn('w:bookmarkStart')}")]
    for needed in ("_Ref_overview", "_Ref_plan", "_Ref_data", "_Ref_ref"):
        assert needed in bm_names, f"북마크 {needed} 누락"

    text = "\n".join(p.text for p in out_doc.paragraphs)
    for needle in ("표준화 종합 보고서", "정보전략실", "단기 과제", "참고"):
        assert needle in text
