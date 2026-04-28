"""30 페이지 분량 데모 SOP `.docx` 결정론적 빌드.

실행:
    cd backend && python -m scripts.build_demo_sop

산출물: backend/app/templates_seed/demo/sop_30p.docx
"""

import io
import struct
import sys
import zipfile
import zlib
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Cm, Inches, Pt

# 프로젝트 루트(backend/) 기준 산출물 경로.
_BACKEND_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = _BACKEND_ROOT / "app" / "templates_seed" / "demo" / "sop_30p.docx"

# 결정성을 위해 zip entry timestamp 를 고정 (1980-01-01 00:00:00 = zip epoch).
_FIXED_ZIP_DT = (1980, 1, 1, 0, 0, 0)
_FIXED_DOC_DT = datetime(2026, 1, 1, 0, 0, 0)


# 1×1 단색 PNG — 결정성을 위해 raw bytes 고정.
def _make_solid_png(rgb: tuple[int, int, int]) -> bytes:
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = b"IHDR" + struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    ihdr_chunk = struct.pack(">I", len(ihdr) - 4) + ihdr + struct.pack(">I", zlib.crc32(ihdr))
    raw = b"\x00" + bytes(rgb)
    idat_data = zlib.compress(raw, 9)
    idat = b"IDAT" + idat_data
    idat_chunk = struct.pack(">I", len(idat) - 4) + idat + struct.pack(">I", zlib.crc32(idat))
    iend = b"IEND"
    iend_chunk = struct.pack(">I", 0) + iend + struct.pack(">I", zlib.crc32(iend))
    return sig + ihdr_chunk + idat_chunk + iend_chunk


_PNG_GRAY = _make_solid_png((180, 180, 180))
_PNG_BLUE = _make_solid_png((90, 130, 200))
_PNG_GREEN = _make_solid_png((130, 180, 130))
_IMAGES = [_PNG_GRAY, _PNG_BLUE, _PNG_GREEN]

# 헤딩 트리 — 4단계까지. 결정적 — 동일 코드, 동일 산출.
_CHAPTERS = [
    ("1. 목적과 적용 범위", [
        ("1.1 목적", []),
        ("1.2 적용 범위", [
            ("1.2.1 포함되는 업무", []),
            ("1.2.2 제외되는 업무", []),
        ]),
    ]),
    ("2. 정의와 약어", [
        ("2.1 용어 정의", [
            ("2.1.1 일반 용어", [("2.1.1.1 핵심 약어", [])]),
            ("2.1.2 시스템 용어", []),
        ]),
        ("2.2 참조 문서", []),
    ]),
    ("3. 책임과 권한", [
        ("3.1 부서별 책임", [
            ("3.1.1 운영팀", []),
            ("3.1.2 품질팀", [("3.1.2.1 검토 절차", [])]),
        ]),
        ("3.2 권한 위임", []),
        ("3.3 비상 권한", []),
    ]),
    ("4. 절차 상세", [
        ("4.1 사전 준비", []),
        ("4.2 실행 단계", [
            ("4.2.1 1 단계 — 입력 검증", []),
            ("4.2.2 2 단계 — 처리", [("4.2.2.1 예외 처리", [])]),
            ("4.2.3 3 단계 — 산출물 검수", []),
        ]),
        ("4.3 사후 검토", []),
    ]),
    ("5. 기록 관리", [
        ("5.1 보존 기간", []),
        ("5.2 폐기 절차", []),
    ]),
]

_LOREM = (
    "본 절차는 표준화된 워크플로우를 정의하며 모든 관련 업무에 일관되게 적용된다. "
    "수행자는 각 단계에서 입력·산출물·책임자를 명확히 식별하고, 이상 발견 시 즉시 보고한다. "
    "본 문서는 분기마다 검토되며, 조직 변경·법규 변경·시스템 변경 시 즉시 개정한다. "
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor "
    "incididunt ut labore et dolore magna aliqua."
)

# landscape 섹션 안에 들어갈 가로 표 (6 컬럼 × 6 행).
_LANDSCAPE_TABLE = [
    ["단계", "담당", "입력", "처리", "산출물", "검토"],
    ["1", "운영팀", "신청서", "검증", "검증 보고", "품질팀"],
    ["2", "품질팀", "검증 보고", "분석", "분석 결과", "관리책임자"],
    ["3", "운영팀", "분석 결과", "조치", "조치 기록", "품질팀"],
    ["4", "관리팀", "조치 기록", "승인", "승인 문서", "경영자"],
    ["5", "전사", "승인 문서", "배포", "배포 확인", "내부감사"],
]

# 본문 표 4개. 일부만 캡션.
_TABLES = [
    {
        "caption": "표 1. 항목별 책임자",
        "rows": [
            ["항목", "주관", "지원"],
            ["입력 검증", "운영팀", "품질팀"],
            ["처리", "운영팀", "시스템팀"],
            ["검수", "품질팀", "관리팀"],
        ],
    },
    {
        "caption": None,  # 캡션 누락
        "rows": [
            ["문서 ID", "버전", "제정일"],
            ["SOP-001", "v1.0", "2026-01-15"],
            ["SOP-002", "v2.1", "2026-03-02"],
        ],
    },
    {
        "caption": "표 3. 검토 주기",
        "rows": [
            ["문서 종류", "검토 주기", "승인자"],
            ["1급", "분기", "경영자"],
            ["2급", "반기", "관리책임자"],
            ["3급", "연간", "팀장"],
        ],
    },
    {
        "caption": None,  # 캡션 누락
        "rows": [
            ["KPI", "목표", "측정 주기"],
            ["불량률", "<0.5%", "월간"],
            ["처리 시간", "<24h", "주간"],
        ],
    },
]


def _add_image(doc: DocxDocument, png_bytes: bytes, *, caption: str | None) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Inches(2.0))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.runs[0].font.size = Pt(9)


def _add_table(doc: DocxDocument, rows: list[list[str]], *, caption: str | None) -> None:
    if caption:
        cap = doc.add_paragraph(caption)
        cap.runs[0].font.size = Pt(9)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            table.rows[ri].cells[ci].text = cell_text


def _add_paragraph_block(doc: DocxDocument, lines: int) -> None:
    for _ in range(lines):
        doc.add_paragraph(_LOREM)


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    new_w, new_h = section.page_height, section.page_width
    section.page_width = new_w
    section.page_height = new_h
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)


def _set_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    new_w, new_h = section.page_height, section.page_width
    section.page_width = new_w
    section.page_height = new_h


def _add_header(section, text: str) -> None:
    para = section.header.paragraphs[0]
    para.text = text


def _build_document() -> DocxDocument:
    doc = Document()

    section = doc.sections[0]
    _add_header(section, "Demo SOP — Word Templator 시연용")

    doc.add_heading("표준 운영 절차 (Demo SOP)", level=0)
    doc.add_paragraph("문서 번호: SOP-DEMO-001 / 개정일: 2026-01-01 / 작성: Word Templator")
    _add_paragraph_block(doc, 2)

    table_iter = iter(_TABLES)
    image_iter = iter(_IMAGES)
    image_caption_iter = iter(["그림 1. 프로세스 흐름", None, None])
    landscape_inserted = False

    def _emit_heading_subtree(items, table_iter, image_iter, image_caption_iter):
        nonlocal landscape_inserted
        for title_text, children in items:
            # 첫 토큰의 비어있지 않은 dot-segment 개수가 헤딩 깊이.
            # "1." → ["1"] → 1, "1.1" → ["1","1"] → 2, "1.2.1" → 3, "2.1.1.1" → 4.
            first = title_text.split()[0]
            level = len([s for s in first.split(".") if s])
            doc.add_heading(title_text, level=min(level, 4))
            _add_paragraph_block(doc, 4)
            if level == 1:
                try:
                    t = next(table_iter)
                    _add_table(doc, t["rows"], caption=t["caption"])
                    _add_paragraph_block(doc, 2)
                except StopIteration:
                    pass
                try:
                    img = next(image_iter)
                    cap = next(image_caption_iter)
                    _add_image(doc, img, caption=cap)
                    _add_paragraph_block(doc, 2)
                except StopIteration:
                    pass
            if level == 1 and title_text.startswith("3.") and not landscape_inserted:
                landscape_section = doc.add_section(WD_SECTION.NEW_PAGE)
                _set_landscape(landscape_section)
                _add_header(landscape_section, "Demo SOP — Word Templator 시연용")
                doc.add_paragraph("부록 A. 가로 페이지 — 단계별 책임 매트릭스").runs[0].font.size = Pt(11)
                _add_table(doc, _LANDSCAPE_TABLE, caption=None)
                _add_paragraph_block(doc, 2)
                portrait_section = doc.add_section(WD_SECTION.NEW_PAGE)
                _set_portrait(portrait_section)
                _add_header(portrait_section, "Demo SOP — Word Templator 시연용")
                landscape_inserted = True
            _emit_heading_subtree(children, table_iter, image_iter, image_caption_iter)

    _emit_heading_subtree(_CHAPTERS, table_iter, image_iter, image_caption_iter)

    # H1 은 1~5장만 — 부록은 H2 로 (test_h1_count_is_5 = 5).
    doc.add_heading("부록 B. 추가 참고", level=2)
    _add_paragraph_block(doc, 30)

    doc.core_properties.author = "Word Templator"
    doc.core_properties.title = "Demo SOP"
    doc.core_properties.created = _FIXED_DOC_DT
    doc.core_properties.modified = _FIXED_DOC_DT
    doc.core_properties.last_modified_by = ""
    return doc


def _normalize_zip_bytes(blob: bytes) -> bytes:
    """zip entry timestamp / external_attr 고정해 결정성 확보."""
    src = zipfile.ZipFile(io.BytesIO(blob))
    out = io.BytesIO()
    dst = zipfile.ZipFile(out, "w", compression=zipfile.ZIP_DEFLATED)
    try:
        for info in sorted(src.infolist(), key=lambda i: i.filename):
            data = src.read(info.filename)
            new_info = zipfile.ZipInfo(filename=info.filename, date_time=_FIXED_ZIP_DT)
            new_info.compress_type = zipfile.ZIP_DEFLATED
            new_info.external_attr = info.external_attr
            dst.writestr(new_info, data)
    finally:
        dst.close()
        src.close()
    return out.getvalue()


def build_demo_sop_bytes() -> bytes:
    """결정론적으로 동일한 .docx 바이트를 반환."""
    doc = _build_document()
    buf = io.BytesIO()
    doc.save(buf)
    return _normalize_zip_bytes(buf.getvalue())


def main() -> int:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PATH.write_bytes(build_demo_sop_bytes())
    print(f"wrote {OUTPUT_PATH} ({OUTPUT_PATH.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
