"""Outline + StyleSpec → .docx 바이트.

섹션(SectionSpec) 단위로 페이지 방향/크기/마진/머리말·꼬리말을 보존한다.
- outline.sections 가 비어있으면 단일 섹션 (legacy / StyleSpec.page 마진 사용).
- 섹션 머리말/꼬리말은 파싱 시 디스크에 저장된 원본 XML 을 다시 주입.
"""

import io
import re
import uuid

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Mm
from lxml import etree

from app.domain.outline import Block, Outline
from app.domain.section import SectionSpec
from app.domain.style_spec import StyleSpec
from app.renderer.apply_style import apply_paragraph_style
from app.renderer.inject_caption_fields import build_caption_paragraph_xml, build_ref_run_xml
from app.renderer.inject_numbering import renumber
from app.renderer.reembed_raw import reembed_paragraph, reembed_table
from app.storage.files import section_part_path


def _default_section_spec(block_ids: list[str], spec: StyleSpec) -> SectionSpec:
    """outline.sections 가 비어있을 때 (legacy) 사용할 단일 portrait 섹션."""
    return SectionSpec(
        id="s-default",
        orientation="portrait",
        margin_top_mm=spec.page.margin_top_mm,
        margin_bottom_mm=spec.page.margin_bottom_mm,
        margin_left_mm=spec.page.margin_left_mm,
        margin_right_mm=spec.page.margin_right_mm,
        block_ids=block_ids,
    )


def _apply_section_props(section, sect_spec: SectionSpec) -> None:
    """python-docx Section 객체에 SectionSpec 의 페이지 속성 적용.

    landscape 시 width/height 를 명시적으로 swap 해야 한다 — python-docx 가 자동
    처리하지 않음.
    """
    is_landscape = sect_spec.orientation == "landscape"
    section.orientation = WD_ORIENT.LANDSCAPE if is_landscape else WD_ORIENT.PORTRAIT
    if is_landscape and sect_spec.page_width_mm < sect_spec.page_height_mm:
        section.page_width = Mm(sect_spec.page_height_mm)
        section.page_height = Mm(sect_spec.page_width_mm)
    else:
        section.page_width = Mm(sect_spec.page_width_mm)
        section.page_height = Mm(sect_spec.page_height_mm)
    section.top_margin = Mm(sect_spec.margin_top_mm)
    section.bottom_margin = Mm(sect_spec.margin_bottom_mm)
    section.left_margin = Mm(sect_spec.margin_left_mm)
    section.right_margin = Mm(sect_spec.margin_right_mm)


def _apply_preserved_part(part_element, xml_bytes: bytes) -> None:
    """파트의 root element 자식들을 보존된 XML 의 자식으로 교체.

    `<w:hdr>` / `<w:ftr>` 의 콘텐츠 (paragraph 들) 만 갈아끼우는 식. 임베디드 이미지의
    rels 까지는 넘기지 않으므로 텍스트 위주 머리말/꼬리말에 한해 안전.
    """
    new_root = etree.fromstring(xml_bytes)
    for child in list(part_element):
        part_element.remove(child)
    for child in list(new_root):
        part_element.append(child)


def _apply_preserved_headers_footers(
    section,
    sect_spec: SectionSpec,
    section_idx: int,
    user_id: uuid.UUID,
    job_id: uuid.UUID,
) -> None:
    """파싱 시 저장된 머리말/꼬리말 XML 을 출력 .docx 의 해당 섹션에 다시 주입."""
    has_first = bool(sect_spec.header_first_ref or sect_spec.footer_first_ref)
    has_even = bool(sect_spec.header_even_ref or sect_spec.footer_even_ref)
    if has_first:
        section.different_first_page_header_footer = True

    targets = [
        ("header", "default", sect_spec.header_default_ref, section.header),
        ("header", "first", sect_spec.header_first_ref, section.first_page_header),
        ("header", "even", sect_spec.header_even_ref, section.even_page_header),
        ("footer", "default", sect_spec.footer_default_ref, section.footer),
        ("footer", "first", sect_spec.footer_first_ref, section.first_page_footer),
        ("footer", "even", sect_spec.footer_even_ref, section.even_page_footer),
    ]
    for kind, position, ref, target in targets:
        if not ref:
            continue
        path = section_part_path(user_id, job_id, section_idx, kind, position)
        if not path.exists():
            continue
        target.is_linked_to_previous = False
        _apply_preserved_part(target.part.element, path.read_bytes())

    # python-docx 가 even/odd 헤더 사용을 settings.xml 에 켜주지 않으므로, even 이
    # 있으면 evenAndOddHeaders 옵션을 직접 설정.
    if has_even:
        try:
            settings = section.part.package.parts
            for p in settings:
                if p.partname.endswith("/word/settings.xml"):
                    settings_root = p.element
                    qn_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                    eaoh = settings_root.find(f"{qn_w}evenAndOddHeaders")
                    if eaoh is None:
                        etree.SubElement(settings_root, f"{qn_w}evenAndOddHeaders")
                    break
        except Exception:
            # 옵션 set 실패해도 본문 렌더는 안 막음.
            pass


_CAPTION_PREFIX_RE = re.compile(
    r"^\s*(그림|표|Figure|Table)\s*(\d+)\s*([.:\]\-—]?)\s*(.*)$"
)


def _add_paragraph_block(doc, block: Block, spec: StyleSpec) -> None:
    """단락 블록 emit — caption_refs 가 있으면 REF 필드로 inline 치환."""
    text = block.text or ""
    if not block.caption_refs:
        para = doc.add_paragraph(text)
        apply_paragraph_style(para, block.level, spec, alignment_override=block.alignment)
        return

    para = doc.add_paragraph()
    apply_paragraph_style(para, block.level, spec, alignment_override=block.alignment)

    cursor = 0
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for ref in sorted(block.caption_refs, key=lambda r: r.span[0]):
        start, end = ref.span
        if start < cursor:  # overlapping match — skip
            continue
        if start > cursor:
            para.add_run(text[cursor:start])
        if ref.target_block_id is None:
            para.add_run(text[start:end])
        else:
            matched_text = text[start:end]
            num_str = str(ref.detected_number)
            # prefix_text: 숫자 앞 레이블 텍스트 (split 첫 조각)
            prefix_text = (
                matched_text.split(num_str)[0] if num_str in matched_text else matched_text
            )
            run_xml = build_ref_run_xml(
                label_kind=ref.label_kind,
                block_id=ref.target_block_id,
                cached_number=ref.detected_number,
                prefix_text=prefix_text,
            )
            wrapper = (
                f'<root xmlns:w="{w_ns}">'.encode()
                + run_xml
                + b"</root>"
            )
            container = etree.fromstring(wrapper)
            for child in list(container):
                para._p.append(child)
        cursor = end
    if cursor < len(text):
        para.add_run(text[cursor:])


def _add_image_placeholder(doc, block: Block, spec: StyleSpec) -> None:
    text = "[이미지]"
    if block.caption:
        text = f"{text} {block.caption}"
    para = doc.add_paragraph(text)
    apply_paragraph_style(para, 0, spec)


def _add_field_placeholder(doc, block: Block, spec: StyleSpec) -> None:
    text = "[참조 — Phase 4 예정]"
    if block.preview_text:
        text = f"{text} {block.preview_text}"
    para = doc.add_paragraph(text)
    apply_paragraph_style(para, 0, spec)


def _add_caption_with_seq(doc, caption: str, block_id: str, spec: StyleSpec) -> None:
    """`그림 N. 제목` / `표 N. 제목` 패턴이면 SEQ 필드로, 아니면 평문 fallback.

    python-docx 의 body.append() 는 <w:sectPr> 뒤에 삽입하므로 <w:sectPr> 앞에
    삽입하기 위해 insert(-1, ...) 를 사용한다.
    """
    m = _CAPTION_PREFIX_RE.match(caption)
    if not m:
        para = doc.add_paragraph(caption)
        apply_paragraph_style(para, 0, spec)
        return
    label, num_str, sep, tail = m.group(1), m.group(2), m.group(3) or ".", m.group(4)
    seq_kind = "Figure" if label in ("그림", "Figure") else "Table"
    xml = build_caption_paragraph_xml(
        label=label,
        seq_kind=seq_kind,
        block_id=block_id,
        cached_number=int(num_str),
        tail_text=f"{sep} {tail}" if tail else "",
    )
    new_p = etree.fromstring(xml)
    # <w:sectPr> 가 body 마지막 자식이므로 그 앞에 삽입 (insert 음수 인덱스 불가 → len-1)
    body = doc.element.body
    body.insert(len(body) - 1, new_p)


def _emit_block(
    doc,
    block: Block,
    spec: StyleSpec,
    *,
    user_id: uuid.UUID | None,
    job_id: uuid.UUID | None,
) -> None:
    if block.kind == "paragraph":
        if block.raw_xml_ref and user_id is not None and job_id is not None:
            reembed_paragraph(doc, raw_ref=block.raw_xml_ref, user_id=user_id, job_id=job_id)
        else:
            _add_paragraph_block(doc, block, spec)
        return

    if block.caption:
        _add_caption_with_seq(doc, block.caption, block.id, spec)

    if block.kind == "table":
        if block.raw_ref and user_id is not None and job_id is not None:
            reembed_table(doc, raw_ref=block.raw_ref, user_id=user_id, job_id=job_id, spec=spec)
        else:
            # raw_ref 없을 때도 실제 Table 요소로 emit — 재파싱 시 Table 탐지 가능.
            # markdown 첫 행에서 컬럼 수를 추정하고 단순 1행 테이블로 근사.
            md = block.markdown or ""
            cols = max(1, md.count("|") // 2) if "|" in md else 1
            doc.add_table(rows=1, cols=cols)
        return

    if block.kind == "image":
        _add_image_placeholder(doc, block, spec)
        return

    if block.kind == "field":
        _add_field_placeholder(doc, block, spec)
        return


def render_docx(
    outline: Outline,
    spec: StyleSpec,
    *,
    user_id: uuid.UUID | None = None,
    job_id: uuid.UUID | None = None,
) -> bytes:
    doc = Document()
    blocks = renumber(outline.blocks, spec)
    blocks_by_id = {b.id: b for b in blocks}

    sections = outline.sections or [_default_section_spec([b.id for b in blocks], spec)]

    for s_idx, sect_spec in enumerate(sections):
        if s_idx == 0:
            section = doc.sections[0]
        else:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
        _apply_section_props(section, sect_spec)
        if user_id is not None and job_id is not None:
            _apply_preserved_headers_footers(section, sect_spec, s_idx, user_id, job_id)

        for block_id in sect_spec.block_ids:
            block = blocks_by_id.get(block_id)
            if block is None:
                continue
            _emit_block(doc, block, spec, user_id=user_id, job_id=job_id)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
