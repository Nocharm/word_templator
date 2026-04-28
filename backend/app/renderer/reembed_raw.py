"""디스크에 저장된 OOXML 조각을 새 .docx body 에 재삽입."""

import uuid

from docx.document import Document as DocxDocument
from lxml import etree

from app.domain.style_spec import StyleSpec
from app.renderer.apply_table_style import apply_table_style
from app.storage.files import raw_ooxml_path


def _parse_fragment(xml_bytes: bytes):
    parser = etree.XMLParser(remove_blank_text=False)
    wrapped = (
        b'<root xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        + xml_bytes
        + b"</root>"
    )
    root = etree.fromstring(wrapped, parser=parser)
    return root[0]


def reembed_table(
    doc: DocxDocument,
    *,
    raw_ref: str,
    user_id: uuid.UUID,
    job_id: uuid.UUID,
    spec: StyleSpec,
) -> None:
    """`raw_ref` 로 디스크에서 <w:tbl> 조각을 읽어 본문 끝에 추가하고 스타일을 덮어쓴다."""
    p = raw_ooxml_path(user_id, job_id, raw_ref)
    if not p.exists():
        doc.add_paragraph(f"[표 원본 누락 — {raw_ref}]")
        return
    try:
        tbl_el = _parse_fragment(p.read_bytes())
        apply_table_style(tbl_el, spec)
        doc.element.body.append(tbl_el)
        doc.add_paragraph()  # python-docx 가 표 직후 빈 문단을 요구하는 경우 안전 패딩
    except (etree.XMLSyntaxError, IndexError):
        # 조각이 손상(잘림/구문 오류/빈 루트)되어 있어도 500 대신 자리표시자로 폴백
        doc.add_paragraph(f"[표 원본 누락 — {raw_ref}]")


def reembed_paragraph(
    doc: DocxDocument,
    *,
    raw_ref: str,
    user_id: uuid.UUID,
    job_id: uuid.UUID,
) -> None:
    """디스크의 원본 <w:p> 를 본문에 그대로 삽입.

    runs / 필드 / 북마크 / 기존 스타일을 모두 보존한다 (Phase 4 trade-off:
    새 StyleSpec 의 폰트/사이즈는 보존된 문단에는 적용되지 않음).
    """
    p = raw_ooxml_path(user_id, job_id, raw_ref)
    if not p.exists():
        doc.add_paragraph(f"[원본 누락 — {raw_ref}]")
        return
    try:
        p_el = _parse_fragment(p.read_bytes())
        doc.element.body.append(p_el)
    except (etree.XMLSyntaxError, IndexError):
        doc.add_paragraph(f"[원본 누락 — {raw_ref}]")
